"""
═══════════════════════════════════════════════════════════
  AI Chat Demo — Advanced RAG Server v2
  ระบบ AI Chat อัจฉริยะ วิเคราะห์เอกสารด้วย Advanced RAG
  Upload → Extract → Chunk → Enrich → Embed → Hybrid Search → AI Answer
═══════════════════════════════════════════════════════════
"""

import os, re, uuid, math, json, hashlib, time, threading, traceback
from concurrent.futures import ThreadPoolExecutor, as_completed
import requests as http_req
from datetime import datetime
import numpy as np
import sqlite3
import faiss
import sqlite3
import faiss

from dotenv import load_dotenv
load_dotenv()

from flask import Flask, request, jsonify, send_from_directory, g, send_file
from flask_cors import CORS
import secrets, io, zipfile

# ═══════════════════════════════════════════════════════
#  Configuration
# ═══════════════════════════════════════════════════════

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
TEXT_CACHE = os.path.join(UPLOAD_FOLDER, '_text_cache')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(TEXT_CACHE, exist_ok=True)

DB_PATH = os.path.join(UPLOAD_FOLDER, 'rag.db')
FAISS_PATH = os.path.join(UPLOAD_FOLDER, 'faiss.index')
FAISS_MAP_PATH = os.path.join(UPLOAD_FOLDER, 'faiss_map.json')
EMBED_DIM = 3072

DB_PATH = os.path.join(UPLOAD_FOLDER, 'rag.db')
FAISS_PATH = os.path.join(UPLOAD_FOLDER, 'faiss.index')
FAISS_MAP_PATH = os.path.join(UPLOAD_FOLDER, 'faiss_map.json')
EMBED_DIM = 3072

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
ALLOWED_EXT = {'.pdf', '.docx', '.doc', '.txt', '.csv', '.xlsx', '.xls', '.md', '.json'}

GEMINI_CONTENT_MODEL = 'gemini-2.5-flash'
GEMINI_CHAT_MODEL = 'gemini-2.5-flash'  # Full model for accurate chat
GEMINI_OCR_MODEL = 'gemini-2.5-flash'  # Vision OCR
GEMINI_EMBED_MODEL = 'gemini-embedding-001'

app = Flask(__name__, static_folder=BASE_DIR, static_url_path='')
CORS(app)

# ═══════════════════════════════════════════════════════
#  Rate Limiter for Gemini Keys
# ═══════════════════════════════════════════════════════

GEMINI_KEYS = []
for k in ['GEMINI_API_KEY', 'GEMINI_API_KEY_2', 'GEMINI_API_KEY_3', 'GEMINI_API_KEY_4', 'GEMINI_API_KEY_5']:
    v = os.environ.get(k, '').strip()
    if v:
        GEMINI_KEYS.append(v)

# Per-key rate tracking
key_usage = {}  # key -> {'calls': [...timestamps], 'cooldown_until': float}
RPM_CONTENT = 25
RPM_EMBED = 1400
COOLDOWN_SECONDS = 2
SKIP_ENRICHMENT = False  # Enable enrichment for better search quality

def _get_best_key(is_embed=False):
    """Pick the Gemini key with lowest recent usage."""
    now = time.time()
    rpm = RPM_EMBED if is_embed else RPM_CONTENT
    best_key, best_count = None, 999999
    for k in GEMINI_KEYS:
        info = key_usage.setdefault(k, {'calls': [], 'cooldown_until': 0})
        if now < info['cooldown_until']:
            continue
        # Count calls in last 60s
        recent = [t for t in info['calls'] if now - t < 60]
        info['calls'] = recent
        if len(recent) < rpm and len(recent) < best_count:
            best_key, best_count = k, len(recent)
    return best_key

def _mark_key_used(key):
    info = key_usage.setdefault(key, {'calls': [], 'cooldown_until': 0})
    info['calls'].append(time.time())

def _mark_key_cooldown(key):
    info = key_usage.setdefault(key, {'calls': [], 'cooldown_until': 0})
    info['cooldown_until'] = time.time() + COOLDOWN_SECONDS
    pass  # Silent cooldown

def _gemini_call(prompt, is_embed=False, max_retries=5, model_name=None):
    """Make a Gemini API call with smart key rotation and rate limiting."""
    import google.generativeai as genai
    use_model = model_name or GEMINI_CONTENT_MODEL
    
    for attempt in range(max_retries):
        key = _get_best_key(is_embed=is_embed)
        if not key:
            if attempt < max_retries - 1:
                wait = min(3, attempt + 1)
                print(f'  ⏳ All keys busy, waiting {wait}s... (attempt {attempt+1}/{max_retries})', flush=True)
                time.sleep(wait)
                continue
            raise Exception('All Gemini API keys exhausted or rate limited')
        
        try:
            genai.configure(api_key=key)
            model = genai.GenerativeModel(use_model)
            _mark_key_used(key)
            response = model.generate_content(prompt)
            return response.text
        except Exception as e:
            err = str(e).lower()
            if '429' in err or 'resource' in err or 'rate' in err:
                _mark_key_cooldown(key)
                if attempt < max_retries - 1:
                    time.sleep(1)
                    continue
            elif '404' in err or 'not found' in err:
                print(f'  ❌ Key {key[:12]}... model not found, trying next', flush=True)
                _mark_key_cooldown(key)
                continue
            raise
    raise Exception('Failed after all retries')


# ═══════════════════════════════════════════════════════
#  In-memory Document Store
# ═══════════════════════════════════════════════════════

documents = []
chat_history = []
enriched_store = {}  # doc_id -> {chunks: [{text, summary, questions, embedding}], docType: {}}

# Scalable storage (SQLite + FAISS)
_db_lock = threading.Lock()
_faiss_lock = threading.Lock()
faiss_index = None   # FAISS vector index
faiss_id_map = []    # FAISS internal idx -> {doc_id, chunk_index}

# Scalable storage (SQLite + FAISS)
_db_lock = threading.Lock()
_faiss_lock = threading.Lock()
faiss_index = None   # FAISS vector index
faiss_id_map = []    # FAISS internal idx -> {doc_id, chunk_index}
upload_progress = {}  # doc_id -> {status, progress, message, steps[], timing{}}

def _update_progress(doc_id, status, progress, message, step_name=None, step_status=None, step_detail=None):
    """Update upload progress with per-step tracking."""
    if doc_id not in upload_progress:
        upload_progress[doc_id] = {
            'status': status, 'progress': progress, 'message': message,
            'steps': [], 'timing': {}, 'startTime': time.time()
        }
    p = upload_progress[doc_id]
    p['status'] = status
    p['progress'] = progress
    p['message'] = message
    
    if step_name:
        # Find existing step or add new
        existing = next((s for s in p['steps'] if s['name'] == step_name), None)
        if existing:
            existing['status'] = step_status or existing['status']
            if step_detail:
                existing['detail'] = step_detail
            if step_status == 'done' and 'startedAt' in existing:
                existing['duration'] = round(time.time() - existing['startedAt'], 1)
        else:
            step = {'name': step_name, 'status': step_status or 'pending', 'detail': step_detail or ''}
            if step_status == 'active':
                step['startedAt'] = time.time()
            p['steps'].append(step)
    
    # Mark step as active with start time
    if step_name and step_status == 'active':
        for s in p['steps']:
            if s['name'] == step_name and 'startedAt' not in s:
                s['startedAt'] = time.time()


# ═══════════════════════════════════════════════════════
#  Cache: Text + Enriched
# ═══════════════════════════════════════════════════════

def _get_text_cache_path(doc_id):
    return os.path.join(TEXT_CACHE, f'{doc_id}.txt')

def _load_cached_text(doc_id):
    path = _get_text_cache_path(doc_id)
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    return None

def _save_cached_text(doc_id, text):
    try:
        path = _get_text_cache_path(doc_id)
        with open(path, 'w', encoding='utf-8') as f:
            f.write(text)
    except Exception as e:
        print(f'[Cache Save Error] {e}', flush=True)

def _save_enriched_cache(doc_id, chunks, doc_type):
    p = os.path.join(TEXT_CACHE, f'{doc_id}_enriched.json')
    data = {'docType': doc_type, 'chunks': [
        {'text': c.get('text',''), 'section': c.get('section',''),
         'summary': c.get('summary',''), 'questions': c.get('questions',[]),
         'embedding': c.get('embedding',[])}
        for c in chunks
    ]}
    with open(p, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False)

def _load_enriched_cache(doc_id):
    p = os.path.join(TEXT_CACHE, f'{doc_id}_enriched.json')
    if os.path.exists(p):
        try:
            with open(p, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            pass
    return None



# ═══════════════════════════════════════════════════════
#  Scalable Storage: SQLite + FAISS
#  Supports thousands of documents efficiently
# ═══════════════════════════════════════════════════════

def _get_db():
    """Get a thread-safe SQLite connection (WAL mode)."""
    conn = sqlite3.connect(DB_PATH, timeout=30)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA synchronous=NORMAL")
    return conn

def _init_database():
    """Create SQLite tables for chunk storage."""
    conn = _get_db()
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS chunks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            doc_id TEXT NOT NULL,
            chunk_index INTEGER NOT NULL,
            text TEXT NOT NULL,
            summary TEXT DEFAULT '',
            questions TEXT DEFAULT '[]',
            section TEXT DEFAULT '',
            UNIQUE(doc_id, chunk_index)
        );
        CREATE INDEX IF NOT EXISTS idx_chunks_doc ON chunks(doc_id);
    """)
    conn.commit()
    conn.close()
    print(f'  💾 SQLite DB ready: {DB_PATH}', flush=True)

def _db_upsert_chunks(doc_id, chunks):
    """Insert or replace chunks for a document in SQLite."""
    with _db_lock:
        conn = _get_db()
        conn.execute('DELETE FROM chunks WHERE doc_id=?', (doc_id,))
        for i, c in enumerate(chunks):
            conn.execute(
                'INSERT INTO chunks (doc_id, chunk_index, text, summary, questions, section) VALUES (?,?,?,?,?,?)',
                (doc_id, i, c.get('text',''), c.get('summary',''),
                 json.dumps(c.get('questions',[]), ensure_ascii=False),
                 c.get('section',''))
            )
        conn.commit()
        conn.close()

def _db_remove_doc(doc_id):
    """Remove all chunks for a document from SQLite."""
    with _db_lock:
        conn = _get_db()
        conn.execute('DELETE FROM chunks WHERE doc_id=?', (doc_id,))
        conn.commit()
        conn.close()

def _db_get_doc_chunks(doc_id):
    """Get all chunks for a document from SQLite."""
    conn = _get_db()
    rows = conn.execute(
        'SELECT chunk_index, text, summary, questions FROM chunks WHERE doc_id=? ORDER BY chunk_index',
        (doc_id,)
    ).fetchall()
    conn.close()
    return [{'chunk_index': r[0], 'text': r[1], 'summary': r[2],
             'questions': json.loads(r[3] or '[]')} for r in rows]

def _db_get_chunks_for_docs(doc_ids):
    """Get all chunks for multiple documents (batch query)."""
    if not doc_ids:
        return []
    conn = _get_db()
    placeholders = ','.join(['?' for _ in doc_ids])
    rows = conn.execute(
        f'SELECT doc_id, chunk_index, text, summary, questions FROM chunks WHERE doc_id IN ({placeholders}) ORDER BY doc_id, chunk_index',
        list(doc_ids)
    ).fetchall()
    conn.close()
    return [{'doc_id': r[0], 'chunk_index': r[1], 'text': r[2],
             'summary': r[3], 'questions': json.loads(r[4] or '[]')} for r in rows]

def _db_total_chunks():
    """Count total chunks in database."""
    try:
        conn = _get_db()
        count = conn.execute('SELECT COUNT(*) FROM chunks').fetchone()[0]
        conn.close()
        return count
    except:
        return 0

def _db_doc_chunk_count(doc_id):
    """Count chunks for a specific document."""
    conn = _get_db()
    count = conn.execute('SELECT COUNT(*) FROM chunks WHERE doc_id=?', (doc_id,)).fetchone()[0]
    conn.close()
    return count

# ── FAISS Vector Index ──────────────────────────────

def _init_faiss_index():
    """Load or create FAISS index for fast vector search."""
    global faiss_index, faiss_id_map
    if os.path.exists(FAISS_PATH) and os.path.exists(FAISS_MAP_PATH):
        try:
            faiss_index = faiss.read_index(FAISS_PATH)
            with open(FAISS_MAP_PATH, 'r') as f:
                faiss_id_map = json.load(f)
            print(f'  FAISS: Loaded {faiss_index.ntotal:,} vectors, {len(faiss_id_map):,} mappings', flush=True)
            return
        except Exception as e:
            print(f'  FAISS load error: {e}', flush=True)
    faiss_index = faiss.IndexFlatIP(EMBED_DIM)
    faiss_id_map = []
    print(f'  FAISS: Created new empty index (dim={EMBED_DIM})', flush=True)

def _save_faiss():
    """Save FAISS index + ID map to disk."""
    try:
        if faiss_index is not None:
            faiss.write_index(faiss_index, FAISS_PATH)
            with open(FAISS_MAP_PATH, 'w') as f:
                json.dump(faiss_id_map, f)
    except Exception as e:
        print(f'  FAISS save error: {e}', flush=True)

def _faiss_add(doc_id, embeddings):
    """Add document embeddings to FAISS index."""
    global faiss_index, faiss_id_map
    if not embeddings:
        return
    valid_vecs = []
    valid_idxs = []
    for ci, emb in enumerate(embeddings):
        if emb and len(emb) == EMBED_DIM and any(v != 0 for v in emb[:5]):
            valid_vecs.append(emb)
            valid_idxs.append(ci)
    if not valid_vecs:
        return
    vecs = np.array(valid_vecs, dtype=np.float32)
    faiss.normalize_L2(vecs)
    with _faiss_lock:
        faiss_index.add(vecs)
        for ci in valid_idxs:
            faiss_id_map.append({'doc_id': doc_id, 'chunk_index': ci})
        _save_faiss()

def _faiss_remove(doc_id):
    """Remove all vectors for a document. Rebuilds index without them."""
    global faiss_index, faiss_id_map
    with _faiss_lock:
        keep_ids = [i for i, m in enumerate(faiss_id_map) if m['doc_id'] != doc_id]
        if len(keep_ids) == len(faiss_id_map):
            return
        if not keep_ids:
            faiss_index = faiss.IndexFlatIP(EMBED_DIM)
            faiss_id_map = []
        else:
            vecs = np.array([faiss_index.reconstruct(i) for i in keep_ids], dtype=np.float32)
            new_map = [faiss_id_map[i] for i in keep_ids]
            faiss_index = faiss.IndexFlatIP(EMBED_DIM)
            faiss_index.add(vecs)
            faiss_id_map = new_map
        _save_faiss()
        print(f'  FAISS: Removed doc {doc_id}, {faiss_index.ntotal} vectors remain', flush=True)

def _faiss_search(query_embedding, top_k=60):
    """Search FAISS for nearest neighbors."""
    if faiss_index is None or faiss_index.ntotal == 0:
        return []
    q = np.array([query_embedding], dtype=np.float32)
    faiss.normalize_L2(q)
    k = min(top_k, faiss_index.ntotal)
    scores, indices = faiss_index.search(q, k)
    results = []
    for score, idx in zip(scores[0], indices[0]):
        idx = int(idx)
        if idx < 0 or idx >= len(faiss_id_map):
            continue
        meta = faiss_id_map[idx]
        results.append({
            'doc_id': meta['doc_id'],
            'chunk_index': meta['chunk_index'],
            'score': float(score),
        })
    return results

def _migrate_to_db():
    """Migrate existing enriched_store data to SQLite + FAISS."""
    if not enriched_store:
        print(f'  No enriched_store data to migrate', flush=True)
        return

    db_count = _db_total_chunks()
    faiss_count = faiss_index.ntotal if faiss_index else 0
    total_enriched = sum(len(d.get('chunks', [])) for d in enriched_store.values())

    if db_count >= total_enriched and faiss_count > 0:
        print(f'  DB already has {db_count:,} chunks, FAISS has {faiss_count:,} vectors - skip migration', flush=True)
        return

    print(f'  Migrating {len(enriched_store)} docs ({total_enriched:,} chunks) to SQLite + FAISS...', flush=True)
    t0 = time.time()

    for doc_id, data in enriched_store.items():
        chunks = data.get('chunks', [])
        if not chunks:
            continue
        _db_upsert_chunks(doc_id, chunks)
        _faiss_remove(doc_id)
        embeddings = [c.get('embedding', []) for c in chunks]
        _faiss_add(doc_id, embeddings)
        doc_name = next((d['name'] for d in documents if d['id'] == doc_id), doc_id)
        n_emb = sum(1 for e in embeddings if e and len(e) == EMBED_DIM and any(v != 0 for v in e[:5]))
        print(f'    {doc_name}: {len(chunks)} chunks, {n_emb} embeddings -> DB+FAISS', flush=True)

    elapsed = time.time() - t0
    print(f'  Migration done in {elapsed:.1f}s: {_db_total_chunks():,} chunks in SQLite, {faiss_index.ntotal:,} vectors in FAISS', flush=True)




# ═══════════════════════════════════════════════════════
#  Metadata Persistence
# ═══════════════════════════════════════════════════════

METADATA_PATH = os.path.join(UPLOAD_FOLDER, '_metadata.json')

# ─── Session / Owner Isolation ────────────────────────
SESSION_COOKIE = 'aichat_sid'

@app.before_request
def _attach_session():
    sid = request.cookies.get(SESSION_COOKIE)
    if not sid or len(sid) < 16:
        sid = secrets.token_hex(16)
        g._new_sid = sid
    g.sid = sid

@app.after_request
def _persist_session(resp):
    new_sid = getattr(g, '_new_sid', None)
    if new_sid:
        resp.set_cookie(SESSION_COOKIE, new_sid, max_age=60*60*24*365,
                        httponly=True, samesite='Lax', path='/')
    return resp

def _can_see(doc):
    o = doc.get('owner', 'public')
    return o == 'public' or o == getattr(g, 'sid', None)

def _visible_docs():
    return [d for d in documents if _can_see(d)]

def _visible_doc_ids():
    return set(d['id'] for d in _visible_docs())


def _save_metadata():
    try:
        meta = []
        for d in documents:
            meta.append({
                'id': d['id'], 'name': d['name'], 'ext': d['ext'],
                'size': d['size'], 'path': d['path'],
                'uploadedAt': d['uploadedAt'],
                'wordCount': d.get('wordCount', 0),
                'chunkCount': len(d.get('chunks', [])),
                'owner': d.get('owner', 'public'),
            })
        with open(METADATA_PATH, 'w', encoding='utf-8') as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f'[Metadata Save Error] {e}', flush=True)

def _load_metadata():
    global documents, enriched_store
    if not os.path.exists(METADATA_PATH):
        return
    try:
        with open(METADATA_PATH, 'r', encoding='utf-8') as f:
            meta = json.load(f)
        for m in meta:
            if not os.path.exists(m.get('path', '')):
                continue
            text = _load_cached_text(m['id']) or ''
            cks = chunk_text(text)
            documents.append({
                'id': m['id'], 'name': m['name'], 'ext': m['ext'],
                'size': m['size'], 'path': m['path'],
                'uploadedAt': m['uploadedAt'],
                'text': text, 'chunks': cks,
                'wordCount': len(text.split()) if text else 0,
                'owner': m.get('owner', 'public'),
            })
            # Load enriched data with embeddings!
            enriched = _load_enriched_cache(m['id'])
            if enriched:
                enriched_store[m['id']] = enriched
                n_emb = sum(1 for c in enriched.get('chunks',[]) if c.get('embedding') and len(c['embedding']) > 0)
                n_sum = sum(1 for c in enriched.get('chunks',[]) if c.get('summary'))
                print(f'    📊 {m["name"]}: {len(enriched.get("chunks",[]))} chunks, {n_sum} summaries, {n_emb} embeddings', flush=True)
        print(f'  📂 Loaded {len(documents)} documents from metadata', flush=True)
    except Exception as e:
        print(f'[Metadata Load Error] {e}', flush=True)
        traceback.print_exc()


# ═══════════════════════════════════════════════════════
#  Text Extraction
# ═══════════════════════════════════════════════════════

def extract_text(filepath, ext, doc_id=None):
    if doc_id:
        cached = _load_cached_text(doc_id)
        if cached:
            return cached
    try:
        if ext == '.pdf':
            text = _extract_pdf(filepath, doc_id=doc_id)
            if doc_id and text and len(text.strip()) > 50:
                _save_cached_text(doc_id, text)
            return text
        elif ext in ('.docx', '.doc'):
            text = _extract_docx(filepath)
            if doc_id and text:
                _save_cached_text(doc_id, text)
            return text
        elif ext in ('.xlsx', '.xls'):
            text = _extract_excel(filepath)
            if doc_id and text:
                _save_cached_text(doc_id, text)
            return text
        elif ext in ('.txt', '.csv', '.md', '.json'):
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            if doc_id and text:
                _save_cached_text(doc_id, text)
            return text
    except Exception as e:
        return f'[Error extracting: {e}]'
    return ''

def _extract_pdf(filepath, doc_id=None):
    """Extract text from PDF. Best of PyMuPDF + PyPDF2, falls back to OCR."""
    text_fitz = ''
    text_pypdf = ''

    # Try PyMuPDF
    try:
        import fitz
        doc = fitz.open(filepath)
        pages = []
        for page in doc:
            t = page.get_text()
            if t and len(t.strip()) > 10:
                pages.append(t)
        doc.close()
        text_fitz = '\n\n'.join(pages)
    except Exception as e:
        print(f'  [PyMuPDF] {e}', flush=True)

    # Try PyPDF2
    try:
        import PyPDF2
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            pages = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
        text_pypdf = '\n\n'.join(pages)
    except Exception as e:
        print(f'  [PyPDF2] {e}', flush=True)

    # Pick the better extractor
    w_fitz = len(text_fitz.split())
    w_pypdf = len(text_pypdf.split())
    if w_fitz > w_pypdf:
        text = text_fitz
        print(f'  [PDF] fitz={w_fitz} > PyPDF2={w_pypdf} -> using fitz', flush=True)
    else:
        text = text_pypdf
        if w_pypdf > 0:
            print(f'  [PDF] PyPDF2={w_pypdf} >= fitz={w_fitz} -> using PyPDF2', flush=True)

    # Smart decision: text extraction vs OCR
    # Count pages and words-per-page to detect partial scans
    try:
        import fitz as _fitz_check
        _doc_check = _fitz_check.open(filepath)
        n_pages = len(_doc_check)
        _doc_check.close()
    except:
        n_pages = max(1, len(text) // 2000)
    
    words = len(text.split())
    words_per_page = words / max(n_pages, 1)
    
    # If good text extraction (>30 words/page on average), use text
    if words_per_page > 30 and len(text.strip()) > 200:
        print(f'  [PDF] Text layer OK: {words} words, {words_per_page:.0f} words/page', flush=True)
        return text
    
    # Low text density = likely scanned → OCR
    print(f'  [PDF] Low text density: {words} words, {words_per_page:.0f} words/page → OCR', flush=True)
    try:
        ocr_text = _ocr_pdf(filepath, doc_id=doc_id)
        if ocr_text and len(ocr_text.strip()) > len(text.strip()):
            return ocr_text
        return text if text else ocr_text
    except Exception as e:
        return text if text else f'[PDF extraction failed: {e}]'


def _ocr_pdf(filepath, doc_id=None):
    """OCR scanned PDF: Single-page Gemini Vision with strict quality control.
    Uses gemini-2.5-flash (full model) for accuracy. Each page OCR'd individually
    to avoid hallucination. Strict garbage detection per page."""
    from PIL import Image
    import io, base64
    from concurrent.futures import ThreadPoolExecutor, as_completed

    print(f'  🔍 Vision OCR (single-page) for {os.path.basename(filepath)}...', flush=True)
    t0 = time.time()

    # ── Step 1: Extract raw images from PDF using PyMuPDF ──
    import fitz as _fitz
    _pdf = _fitz.open(filepath)
    _raw_images = []
    for _pi in range(len(_pdf)):
        _page = _pdf[_pi]
        _imgs = _page.get_images()
        if _imgs:
            _xref = _imgs[0][0]
            _bi = _pdf.extract_image(_xref)
            _raw_images.append(_bi["image"])
            del _bi
        else:
            _mat = _fitz.Matrix(1.5, 1.5)  # ~108 dpi
            _pix = _page.get_pixmap(matrix=_mat)
            _raw_images.append(_pix.tobytes("jpeg"))
            del _pix
    _pdf.close()
    total_pages = len(_raw_images)
    print(f'  📄 {total_pages} pages extracted in {time.time()-t0:.1f}s', flush=True)

    if doc_id:
        _update_progress(doc_id, 'processing', 8, f'Starting OCR for {total_pages} pages...', 'extract', 'active', f'{total_pages} pages')

    # ── Step 2: OCR each page individually with Gemini Vision ──
    MAX_CHARS_PER_PAGE = 8000  # Hallucination threshold per page
    pages_text = {}  # page_num -> text
    pages_lock = threading.Lock()
    done_count = [0]

    def _ocr_single_page(page_idx, raw_img):
        """OCR one page with strict quality control."""
        try:
            img = Image.open(io.BytesIO(raw_img))
            w, h = img.size
            # Resize for optimal OCR (max 1200px wide)
            if w > 1200:
                ratio = 1200 / w
                img = img.resize((1200, int(h * ratio)), Image.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format='JPEG', quality=75)
            b64 = base64.b64encode(buf.getvalue()).decode()
        except Exception as e:
            print(f'    ❌ Page {page_idx+1} image prep: {e}', flush=True)
            return

        prompt = (
            f"อ่านข้อความทั้งหมดจากภาพเอกสารหน้า {page_idx+1} "
            "ให้ครบทุกตัวอักษร ทุกตาราง ทุกรายการ ทุกตัวเลข "
            "สำคัญ: อ่านตัวเลขไทย (๑๒๓๔๕๖๗๘๙๐) และเปอร์เซ็นต์ให้ถูกต้องตามต้นฉบับ "
            "เป็นภาษาไทยและภาษาอังกฤษตามต้นฉบับ ไม่ต้องสรุป ไม่ต้องวิเคราะห์ แค่ถอดข้อความ"
        )

        for attempt in range(3):
            api_key = _get_best_key(is_embed=False)
            if not api_key:
                api_key = GEMINI_KEYS[(page_idx + attempt) % len(GEMINI_KEYS)]
                cd_until = key_usage.get(api_key, {}).get('cooldown_until', 0)
                if cd_until > time.time():
                    time.sleep(min(cd_until - time.time() + 0.2, 3))

            try:
                url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_OCR_MODEL}:generateContent?key={api_key}"
                payload = {
                    "contents": [{"parts": [
                        {"text": prompt},
                        {"inline_data": {"mime_type": "image/jpeg", "data": b64}}
                    ]}],
                    "generationConfig": {"temperature": 0.1, "maxOutputTokens": 8192}
                }
                resp = http_req.post(url, json=payload, timeout=45)
                _mark_key_used(api_key)

                if resp.status_code == 200:
                    data = resp.json()
                    cands = data.get('candidates', [])
                    if cands:
                        text = ''.join(p.get('text', '') for p in cands[0].get('content', {}).get('parts', []))
                        finish = cands[0].get('finishReason', '')

                        # RECITATION — skip this page
                        if finish == 'RECITATION':
                            print(f'    ⚠️ Page {page_idx+1}: RECITATION', flush=True)
                            return

                        # Garbage check: hallucination produces way too many chars
                        if text and len(text) > MAX_CHARS_PER_PAGE:
                            # Check if it's actual content or garbage dots/repeats
                            garbage = sum(1 for c in text if c in '.…_=~-')
                            if garbage / max(len(text), 1) > 0.2:
                                print(f'    🗑️ Page {page_idx+1}: {len(text)} chars, {garbage/len(text):.0%} garbage — retry', flush=True)
                                time.sleep(1)
                                continue
                            # Even without garbage chars, too many chars = hallucination
                            print(f'    🗑️ Page {page_idx+1}: {len(text)} chars (>{MAX_CHARS_PER_PAGE}) — hallucination, retry', flush=True)
                            time.sleep(1)
                            continue

                        if text and len(text.strip()) > 20:
                            with pages_lock:
                                pages_text[page_idx] = text
                                done_count[0] += 1
                                if doc_id:
                                    pct = 10 + int(70 * done_count[0] / total_pages)
                                    _update_progress(doc_id, 'processing', pct,
                                        f'OCR: {done_count[0]}/{total_pages} pages...',
                                        'extract', 'active', f'OCR: {done_count[0]}/{total_pages} pages done')
                            print(f'    ✅ Page {page_idx+1}: {len(text)} chars', flush=True)
                            return

                elif resp.status_code == 429:
                    _mark_key_cooldown(api_key)
                    time.sleep(1 * (attempt + 1))
                    continue

            except Exception as e:
                print(f'    ❌ Page {page_idx+1} attempt {attempt+1}: {e}', flush=True)
                time.sleep(1)

        # All attempts failed
        print(f'    ⚠️ Page {page_idx+1}: all attempts failed', flush=True)

    # ── Run single-page OCR in parallel ──
    n_workers = min(2, len(GEMINI_KEYS), total_pages)  # Match key count
    print(f'  🚀 OCR {total_pages} pages with {n_workers} parallel workers...', flush=True)

    with ThreadPoolExecutor(max_workers=n_workers) as pool:
        futs = []
        for i in range(total_pages):
            f = pool.submit(_ocr_single_page, i, _raw_images[i])
            futs.append(f)
        for f in as_completed(futs):
            try:
                f.result()
            except Exception:
                pass

    del _raw_images  # Free memory

    # ── Assemble final text ──
    all_text = []
    for i in range(total_pages):
        if i in pages_text:
            all_text.append(f'--- Page {i+1} ---\n{pages_text[i]}')

    result = '\n\n'.join(all_text)

    # Clean garbage patterns
    result = re.sub(r'[.…]{20,}', '...', result)
    result = re.sub(r'[-_=~]{20,}', '---', result)
    result = re.sub(r'\n{5,}', '\n\n\n', result)

    elapsed = time.time() - t0
    print(f'  📊 OCR done: {len(result):,} chars, {len(pages_text)}/{total_pages} pages in {elapsed:.1f}s', flush=True)

    if doc_id and doc_id in upload_progress:
        _update_progress(doc_id, 'processing', 80, f'OCR: {len(pages_text)}/{total_pages} pages OK',
                        'extract', 'done', f'{len(pages_text)}/{total_pages} pages, {len(result):,} chars')

    return result


def _extract_docx(filepath):
    """Extract text from Word documents — paragraphs + tables + headers."""
    try:
        from docx import Document
        doc = Document(filepath)
        parts = []
        
        # Extract all paragraphs with style info
        for p in doc.paragraphs:
            txt = p.text.strip()
            if not txt:
                continue
            # Mark headings for section-aware chunking
            style = p.style.name if p.style else ''
            if 'Heading' in style or 'หัวข้อ' in style:
                parts.append(f'\n## {txt}')
            else:
                parts.append(txt)
        
        # Extract all tables
        for ti, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(' | '.join(cells))
            if rows:
                parts.append(f'\n[ตาราง {ti+1}]')
                parts.extend(rows)
        
        return '\n'.join(parts)
    except ImportError:
        return '[กรุณาติดตั้ง python-docx: pip install python-docx]'
    except Exception as e:
        return f'[Error extracting DOCX: {e}]'

def _extract_excel(filepath):
    """Extract text from Excel — all sheets with structured formatting."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
        lines = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            lines.append(f'\n## Sheet: {sheet}')
            header = None
            for ri, row in enumerate(ws.iter_rows(values_only=True)):
                cells = [str(c).strip() if c is not None else '' for c in row]
                # Skip completely empty rows
                if not any(cells):
                    continue
                # First non-empty row = header
                if header is None:
                    header = cells
                    lines.append(' | '.join(header))
                    lines.append(' | '.join(['---'] * len(header)))
                else:
                    # Format as key-value if only 2 columns
                    if len(cells) == 2 and cells[0] and cells[1]:
                        lines.append(f'{cells[0]}: {cells[1]}')
                    else:
                        lines.append(' | '.join(cells))
        wb.close()
        return '\n'.join(lines)
    except ImportError:
        return '[กรุณาติดตั้ง openpyxl: pip install openpyxl]'
    except Exception as e:
        return f'[Error extracting Excel: {e}]'


# ═══════════════════════════════════════════════════════
#  Advanced Text Chunking (RAG)
# ═══════════════════════════════════════════════════════

def chunk_text(text, size=800, overlap=100):
    """Section-aware text chunking. Respects section boundaries (##, ---, Page markers).
    Keeps numbered lists and tables together. Never breaks mid-list."""
    if not text or len(text) < 10:
        return []
    
    # Split into paragraphs
    paragraphs = re.split(r'\n\s*\n|\n', text)
    
    # Detect section boundaries — these are strong break points
    def _is_section_break(para):
        p = para.strip()
        if re.match(r'^---\s*Page\s*\d+', p): return True   # Page markers
        if re.match(r'^#{1,3}\s', p): return True             # Markdown headings
        if re.match(r'^[๐-๙]+\.\s', p): return True          # Thai numbered sections (top-level only, e.g. ๘.)
        if re.match(r'^\d+\.\s+[ก-์A-Z]', p): return True    # Arabic numbered sections
        if re.match(r'^===', p): return True                   # Sheet markers
        return False
    
    def _is_continuation(para):
        """Check if paragraph continues a numbered list or sub-item."""
        p = para.strip()
        if re.match(r'^[๐-๙]+\.[๐-๙]+', p): return True    # Sub-items like ๘.๑
        if re.match(r'^\d+\.\d+', p): return True
        if re.match(r'^[-•]\s', p): return True               # Bullet points
        if re.match(r'^\(\d+\)', p): return True              # (1), (2)
        if re.match(r'^\([ก-์]\)', p): return True            # (ก), (ข)
        return False
    
    chunks, buf, buf_len = [], [], 0
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        # Force break at section boundaries (but only if buffer is non-tiny)
        if _is_section_break(para) and buf and buf_len > 100:
            chunks.append('\n'.join(buf))
            # Overlap: keep last portion
            keep, klen = [], 0
            for b in reversed(buf):
                if klen + len(b) > overlap:
                    break
                keep.insert(0, b)
                klen += len(b)
            buf, buf_len = keep, klen
        
        # Normal size-based break (but don't break continuation items)
        elif buf_len + len(para) > size and buf and not _is_continuation(para):
            chunks.append('\n'.join(buf))
            keep, klen = [], 0
            for b in reversed(buf):
                if klen + len(b) > overlap:
                    break
                keep.insert(0, b)
                klen += len(b)
            buf, buf_len = keep, klen
        
        buf.append(para)
        buf_len += len(para)
    
    if buf:
        chunks.append('\n'.join(buf))
    
    return chunks


# ═══════════════════════════════════════════════════════
#  Enrichment: Summary + Questions via Gemini
# ═══════════════════════════════════════════════════════

def _enrich_chunks_batch(chunks, doc_name, doc_type_info=None, doc_id=None, batch_size=20):
    """Batch-enrich chunks with summaries and questions."""
    total = len(chunks)
    t_enrich = time.time()
    print(f'  ✨ Enriching {total} chunks for {doc_name}...', flush=True)
    consecutive_fails = 0

    for i in range(0, total, batch_size):
        if consecutive_fails >= 3:
            print(f'  ⚠️ Too many failures ({consecutive_fails}), stopping enrichment', flush=True)
            break
        batch = chunks[i:i+batch_size]
        chunk_texts = '\n\n'.join(f'[Chunk {j+1}]\n{c["text"][:500]}' for j, c in enumerate(batch))
        prompt = f"""Analyze these text chunks from document "{doc_name}".
For each chunk, generate: 1. summary (Thai, 1-2 sentences) 2. questions (Thai, 2-3 questions that this chunk can answer)

{chunk_texts}

Reply ONLY valid JSON array: [{{"chunk":1,"summary":"...","questions":["?","?"]}}]"""
        try:
            resp = _gemini_call(prompt, is_embed=False)
            # Parse JSON from response
            resp_clean = resp.strip()
            if '```' in resp_clean:
                resp_clean = re.search(r'```(?:json)?\s*(.*?)```', resp_clean, re.DOTALL)
                resp_clean = resp_clean.group(1) if resp_clean else resp.strip()
            enrichments = json.loads(resp_clean)
            for enr in enrichments:
                idx = enr.get('chunk', 1) - 1
                if 0 <= idx < len(batch):
                    batch[idx]['summary'] = enr.get('summary', '')
                    batch[idx]['questions'] = enr.get('questions', [])
            consecutive_fails = 0
            print(f'    ✅ Enriched batch {i+1}-{min(i+batch_size, total)}/{total}', flush=True)
        except Exception as e:
            consecutive_fails += 1
            print(f'    ❌ Batch {i+1}-{min(i+batch_size, total)} failed: {e}', flush=True)
            time.sleep(1)

    enriched = sum(1 for c in chunks if c.get('summary'))
    t_enrich_end = time.time() - t_enrich
    print(f'  📊 Enriched: {enriched}/{total} chunks — ⏱️ {t_enrich_end:.1f}s', flush=True)
    return chunks


def _enrich_chunks_parallel(chunks, doc_name, doc_id=None, batch_size=30):
    """Parallel enrichment using multiple API keys for speed."""
    total = len(chunks)
    t0 = time.time()
    print(f'  ✨ Parallel enriching {total} chunks for {doc_name} ({len(GEMINI_KEYS)} keys)...', flush=True)

    # Build batches
    batches = []
    for i in range(0, total, batch_size):
        batches.append((i, chunks[i:i+batch_size]))

    completed = [0]
    lock = threading.Lock()

    def _enrich_one_batch(batch_idx, start_idx, batch_chunks, api_key):
        chunk_texts = '\n\n'.join(
            f'[Chunk {j+1}]\n{c["text"][:300]}' for j, c in enumerate(batch_chunks)
        )
        prompt = f"""Analyze these text chunks from document "{doc_name}".
For each chunk, generate: 1. summary (Thai, 1-2 sentences) 2. questions (Thai, 2-3 questions)

{chunk_texts}

Reply ONLY valid JSON array: [{{"chunk":1,"summary":"...","questions":["?","?"]}}]"""

        for attempt in range(5):
            try:
                url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_CONTENT_MODEL}:generateContent?key={api_key}"
                payload = {
                    "contents": [{"parts": [{"text": prompt}]}],
                    "generationConfig": {"temperature": 0.2, "maxOutputTokens": 16384}
                }
                resp = http_req.post(url, json=payload, timeout=60)
                
                if resp.status_code == 429:
                    _mark_key_cooldown(api_key)
                    wait = 1 * (attempt + 1)
                    print(f'    ⏸️ Enrich batch {start_idx+1} key ...{api_key[-6:]} 429, wait {wait}s, rotating...', flush=True)
                    time.sleep(wait)
                    # Rotate key
                    try:
                        ki = GEMINI_KEYS.index(api_key)
                    except ValueError:
                        ki = batch_idx
                    api_key = GEMINI_KEYS[(ki + 1) % len(GEMINI_KEYS)]
                    continue

                if resp.status_code != 200:
                    print(f'    ⚠️ Enrich batch {start_idx+1} HTTP {resp.status_code}', flush=True)
                    if attempt < 4:
                        time.sleep(1)
                    continue

                data = resp.json()
                candidates = data.get('candidates', [])
                if not candidates:
                    continue

                text = ''.join(p.get('text', '') for p in candidates[0].get('content', {}).get('parts', []))
                if not text:
                    continue

                # Parse JSON — robust extraction
                text_raw = text.strip()
                enrichments = None
                # Method 1: find JSON array anywhere in text
                arr_match = re.search(r'\[\s*\{.*\}\s*\]', text_raw, re.DOTALL)
                if arr_match:
                    try:
                        enrichments = json.loads(arr_match.group())
                    except json.JSONDecodeError:
                        pass
                # Method 2: try full text after stripping code fences
                if enrichments is None:
                    text_clean = text_raw
                    if '```' in text_clean:
                        m = re.search(r'```(?:json)?\s*(.*?)```', text_clean, re.DOTALL)
                        if m:
                            text_clean = m.group(1).strip()
                    if text_clean:
                        try:
                            enrichments = json.loads(text_clean)
                        except json.JSONDecodeError:
                            pass
                # Method 3: repair truncated JSON
                if enrichments is None:
                    # Find partial array and try to close it
                    arr_start = text_raw.find('[')
                    if arr_start >= 0:
                        partial = text_raw[arr_start:]
                        # Try closing at last complete object
                        last_brace = partial.rfind('}')
                        if last_brace > 0:
                            repaired = partial[:last_brace+1] + ']'
                            try:
                                enrichments = json.loads(repaired)
                                print(f'    🔧 Enrich batch {start_idx+1} repaired truncated JSON: {len(enrichments)} items', flush=True)
                            except json.JSONDecodeError:
                                pass
                if enrichments is None:
                    print(f'    ⚠️ Enrich batch {start_idx+1} unparseable (len={len(text_raw)}): {text_raw[:200]}', flush=True)
                    if attempt < 4:
                        time.sleep(1)
                    continue
                for enr in enrichments:
                    idx = enr.get('chunk', 1) - 1
                    if 0 <= idx < len(batch_chunks):
                        batch_chunks[idx]['summary'] = enr.get('summary', '')
                        batch_chunks[idx]['questions'] = enr.get('questions', [])

                _mark_key_used(api_key)
                with lock:
                    completed[0] += 1
                    if doc_id and doc_id in upload_progress:
                        pct = 84 + int(1 * completed[0] / len(batches))
                        done_n = completed[0] * batch_size
                        _update_progress(doc_id, 'processing', pct,
                                       f'Enriching {min(done_n, total)}/{total}...',
                                       'enrich', 'active', f'{completed[0]}/{len(batches)} batches done')
                
                print(f'    ✅ Enriched batch {start_idx+1}-{min(start_idx+batch_size, total)}/{total}', flush=True)
                return True

            except json.JSONDecodeError as jde:
                print(f'    ⚠️ Enrich batch {start_idx+1} JSON error: {str(jde)[:80]}', flush=True)
                if attempt < 4:
                    time.sleep(1)
            except Exception as e:
                print(f'    ❌ Enrich batch {start_idx+1}: {e}', flush=True)
                if attempt < 4:
                    time.sleep(1)
        # Fallback: auto-generate minimal summaries from chunk text
        for c in batch_chunks:
            if not c.get('summary'):
                c['summary'] = c['text'][:200].strip() + '...'
                c['questions'] = []
        print(f'    ⚠️ Enrich batch {start_idx+1} FALLBACK: auto-summaries for {len(batch_chunks)} chunks', flush=True)
        with lock:
            completed[0] += 1
        return False

    # Launch parallel enrichment
    if not batches:
        print('  ⚠️ No batches to enrich, returning chunks as-is', flush=True)
        return chunks
    num_workers = max(1, min(4, len(batches), 8))  # Use up to 4 workers
    print(f'  🚀 Enriching {len(batches)} batches with {num_workers} workers...', flush=True)

    with ThreadPoolExecutor(max_workers=num_workers) as pool:
        futures = {}
        for idx, (start, batch) in enumerate(batches):
            key = GEMINI_KEYS[idx % len(GEMINI_KEYS)]
            f = pool.submit(_enrich_one_batch, idx, start, batch, key)
            futures[f] = idx

        for f in as_completed(futures):
            try:
                f.result()
            except Exception as e:
                print(f'    ❌ Enrich exception: {e}', flush=True)

    enriched = sum(1 for c in chunks if c.get('summary'))
    elapsed = time.time() - t0
    print(f'  📊 Parallel enriched: {enriched}/{total} chunks — ⏱️ {elapsed:.1f}s', flush=True)
    return chunks


# ═══════════════════════════════════════════════════════
#  Batch Embedding via Gemini
# ═══════════════════════════════════════════════════════

def _batch_embed_gemini(texts, task_type='RETRIEVAL_DOCUMENT', batch_size=100, doc_id=None, total_chunks=0):
    """Batch embed texts using Gemini's batchEmbedContents API."""
    import google.generativeai as genai
    
    all_embeddings = []
    for start in range(0, len(texts), batch_size):
        batch_texts = texts[start:start+batch_size]
        key = _get_best_key(is_embed=True)
        if not key:
            print(f'  ⚠️ No embed key available, using zeros', flush=True)
            all_embeddings.extend([[0.0]*3072] * len(batch_texts))
            continue
        
        for attempt in range(3):
            try:
                genai.configure(api_key=key)
                _mark_key_used(key)
                result = genai.embed_content(
                    model=f'models/{GEMINI_EMBED_MODEL}',
                    content=batch_texts,
                    task_type=task_type,
                )
                embs = result.get('embedding', [])
                if embs and isinstance(embs[0], list):
                    all_embeddings.extend(embs)
                elif embs:
                    all_embeddings.append(embs)
                print(f'    ✅ Embedded batch {start+1}-{start+len(batch_texts)}/{len(texts)}', flush=True)
                if doc_id and doc_id in upload_progress:
                    done_vecs = start + len(batch_texts)
                    tc = total_chunks or len(texts)
                    e_pct = 85 + int(12 * done_vecs / tc)
                    _update_progress(doc_id, 'processing', e_pct,
                                    f'Embedding {done_vecs}/{tc}...',
                                    'embed', 'active', f'{done_vecs}/{tc} vectors')
                break
            except Exception as e:
                err = str(e).lower()
                if '429' in err or 'rate' in err:
                    _mark_key_cooldown(key)
                    key = _get_best_key(is_embed=True)
                    if not key:
                        print(f'  ⚠️ All embed keys exhausted', flush=True)
                        all_embeddings.extend([[0.0]*3072] * len(batch_texts))
                        break
                    time.sleep(1)
                else:
                    print(f'    ❌ Embed error: {e}', flush=True)
                    all_embeddings.extend([[0.0]*3072] * len(batch_texts))
                    break
    
    return all_embeddings

def _embed_query(text):
    """Embed a single query text."""
    results = _batch_embed_gemini([text], task_type='RETRIEVAL_QUERY', batch_size=1)
    return results[0] if results else [0.0]*3072


# ═══════════════════════════════════════════════════════
#  Pipeline: Extract → Chunk → Enrich → Embed
# ═══════════════════════════════════════════════════════

def _run_pipeline(doc_id, doc_name, text):
    """Run enrichment + embedding pipeline in background."""
    try:
        t0 = time.time()
        print(f'\n  🚀 Pipeline start: {doc_name}', flush=True)
        chunks_raw = chunk_text(text)
        chunks = [{'text': c, 'section': '', 'summary': '', 'questions': [], 'embedding': []} for c in chunks_raw]
        total = len(chunks)
        print(f'  📝 {total} chunks created', flush=True)

        # Step 1: Enrich with summaries
        if total <= 150:
            chunks = _enrich_chunks_batch(chunks, doc_name, doc_id=doc_id)

        # Step 2: Embed
        embed_texts = []
        for c in chunks:
            et = c['text']
            if c.get('summary'):
                et = c['summary'] + '\n' + et
            embed_texts.append(et[:2000])
        
        embeddings = _batch_embed_gemini(embed_texts)
        for j, emb in enumerate(embeddings):
            if j < len(chunks):
                chunks[j]['embedding'] = emb

        # Step 3: Save
        n_emb = sum(1 for c in chunks if c.get('embedding') and len(c['embedding']) > 0 and any(v != 0 for v in c['embedding'][:5]))
        n_sum = sum(1 for c in chunks if c.get('summary'))
        doc_type = {'type': 'general'}
        _save_enriched_cache(doc_id, chunks, doc_type)
        enriched_store[doc_id] = {'docType': doc_type, 'chunks': chunks}
        # Persist to SQLite + FAISS
        _db_upsert_chunks(doc_id, chunks)
        _faiss_remove(doc_id)
        _faiss_add(doc_id, [c.get('embedding', []) for c in chunks])

        # Update document metadata
        for d in documents:
            if d['id'] == doc_id:
                d['chunkCount'] = total
                break
        _save_metadata()
        
        elapsed = time.time() - t0
        print(f'  ✅ Pipeline done: {doc_name} — {total} chunks, {n_sum} summaries, {n_emb} embeddings — ⏱️ {elapsed:.1f}s', flush=True)
    except Exception as e:
        print(f'  ❌ Pipeline error: {e}', flush=True)
        traceback.print_exc()


# ═══════════════════════════════════════════════════════
#  Hybrid Search: Keyword (BM25-like) + Vector + RRF
# ═══════════════════════════════════════════════════════

def _cosine_sim(a, b):
    """Cosine similarity between two vectors."""
    a, b = np.array(a, dtype=np.float32), np.array(b, dtype=np.float32)
    na, nb = np.linalg.norm(a), np.linalg.norm(b)
    if na == 0 or nb == 0:
        return 0.0
    return float(np.dot(a, b) / (na * nb))

def search_docs(query, top_k=15, allowed_doc_ids=None):
    """Hybrid search: FAISS vector + keyword scoring + RRF fusion.
    If allowed_doc_ids is provided, results restricted to that set (multi-user isolation)."""
    if not documents:
        return []

    q_terms = [t for t in query.lower().split() if len(t) > 1]
    if not q_terms:
        return []

    doc_names = {d['id']: d['name'] for d in documents}

    # ── Step 1: FAISS vector search (fast, O(log n)) ──
    vector_hits = {}  # (doc_id, chunk_index) -> score
    candidate_doc_ids = set()

    if faiss_index and faiss_index.ntotal > 0 and GEMINI_KEYS:
        try:
            q_emb = _embed_query(query)
            faiss_results = _faiss_search(q_emb, top_k=60)
            for r in faiss_results:
                if allowed_doc_ids is not None and r['doc_id'] not in allowed_doc_ids:
                    continue
                key = (r['doc_id'], r['chunk_index'])
                vector_hits[key] = r['score']
                candidate_doc_ids.add(r['doc_id'])
        except Exception as e:
            print(f'  ⚠️ Vector search error: {e}', flush=True)

    # ── Step 2: Load candidate chunks from SQLite ──
    if not candidate_doc_ids:
        if allowed_doc_ids is not None:
            candidate_doc_ids = set(allowed_doc_ids)
        else:
            candidate_doc_ids = set(d['id'] for d in documents)
    elif allowed_doc_ids is not None:
        candidate_doc_ids &= set(allowed_doc_ids)

    candidate_chunks = _db_get_chunks_for_docs(candidate_doc_ids)

    if not candidate_chunks:
        # Fallback: try enriched_store (in case DB not populated yet)
        for doc in documents:
            doc_enriched = enriched_store.get(doc['id'], {}).get('chunks', [])
            for ci, ec in enumerate(doc_enriched):
                candidate_chunks.append({
                    'doc_id': doc['id'], 'chunk_index': ci,
                    'text': ec.get('text', ''), 'summary': ec.get('summary', ''),
                    'questions': ec.get('questions', []),
                })

    if not candidate_chunks:
        return []

    # ── Step 3: Keyword scoring on candidates ──
    keyword_hits = {}
    for c in candidate_chunks:
        searchable = (c['text'] + ' ' + c.get('summary', '') + ' ' +
                      ' '.join(c.get('questions', []))).lower()
        score = 0
        matched = []
        for t in q_terms:
            cnt = searchable.count(t)
            if cnt > 0:
                tf = cnt / max(len(searchable.split()), 1)
                idf = math.log(1 + len(candidate_chunks) / max(1, sum(
                    1 for ch in candidate_chunks
                    if t in (ch['text'] + ' ' + ch.get('summary', '')).lower()
                )))
                score += tf * idf * 100
                matched.append(t)
        if matched:
            coverage = len(set(matched)) / len(set(q_terms))
            score *= (1 + coverage)
        if score > 0:
            key = (c['doc_id'], c['chunk_index'])
            keyword_hits[key] = score

    # ── Step 4: RRF Fusion ──
    all_keys = set(vector_hits.keys()) | set(keyword_hits.keys())
    if not all_keys:
        return []

    k_rrf = 60
    vs_sorted = sorted(vector_hits.items(), key=lambda x: x[1], reverse=True)
    vs_rank = {key: rank for rank, (key, _) in enumerate(vs_sorted)}
    kw_sorted = sorted(keyword_hits.items(), key=lambda x: x[1], reverse=True)
    kw_rank = {key: rank for rank, (key, _) in enumerate(kw_sorted)}

    fused = []
    for key in all_keys:
        rrf = 0
        if key in vector_hits:
            rrf += 1.0 / (k_rrf + vs_rank.get(key, len(all_keys)))
        if key in keyword_hits:
            rrf += 1.0 / (k_rrf + kw_rank.get(key, len(all_keys)))
        fused.append({
            'key': key,
            'rrf_score': rrf,
            'vs_score': vector_hits.get(key, 0),
            'kw_score': keyword_hits.get(key, 0),
        })

    fused.sort(key=lambda x: x['rrf_score'], reverse=True)

    # ── Step 5: Build results ──
    chunk_lookup = {(c['doc_id'], c['chunk_index']): c for c in candidate_chunks}
    results = []
    seen = set()
    for f in fused:
        doc_id, ci = f['key']
        sig = f'{doc_id}_{ci}'
        if sig in seen:
            continue
        seen.add(sig)

        chunk_data = chunk_lookup.get((doc_id, ci), {})
        results.append({
            'docId': doc_id,
            'docName': doc_names.get(doc_id, 'Unknown'),
            'chunk': chunk_data.get('text', ''),
            'summary': chunk_data.get('summary', ''),
            'chunkIndex': ci,
            'score': round(f['rrf_score'], 6),
            'kw_score': round(f['kw_score'], 4),
            'vs_score': round(f['vs_score'], 4),
            'matchedTerms': [],
            'coverage': 0,
        })
        if len(results) >= top_k:
            break

    return results


# ═══════════════════════════════════════════════════════
#  Level 4: Query Rewrite + Multi-Query Retrieval
# ═══════════════════════════════════════════════════════

def _rewrite_query(question, history=None):
    """Use AI to rewrite/expand the query for better retrieval."""
    if not GEMINI_KEYS:
        return [question]
    
    hist_ctx = ''
    if history:
        recent = history[-3:]
        hist_ctx = '\nบทสนทนาก่อนหน้า:\n' + '\n'.join(
            f'User: {h.get("question","")}' for h in recent
        )
    
    prompt = f"""คุณเป็นผู้เชี่ยวชาญด้านการค้นหาเอกสาร
งาน: แปลงคำถามของผู้ใช้เป็น 2 search queries ที่แตกต่างกัน เพื่อค้นหาข้อมูลให้ครอบคลุมที่สุด
{hist_ctx}
คำถามเดิม: {question}

กฎ:
1. Query 1 = คำถามเดิม (ปรับให้ชัดเจนขึ้น)
2. Query 2 = มุมมองต่าง เช่น synonym หรือ keyword สำคัญ
3. ทุก query ต้องเป็นภาษาเดียวกับคำถามเดิม

ตอบเป็น JSON array เท่านั้น: ["query1", "query2"]"""
    
    try:
        resp = _gemini_call(prompt, is_embed=False)
        resp_clean = resp.strip()
        if '```' in resp_clean:
            m = re.search(r'```(?:json)?\s*(.*?)```', resp_clean, re.DOTALL)
            if m:
                resp_clean = m.group(1)
        queries = json.loads(resp_clean)
        if isinstance(queries, list) and len(queries) > 0:
            # Always include original question
            result = [question] + [q for q in queries if q != question]
            print(f'  🔄 Query rewrite: {len(result)} queries', flush=True)
            for i, q in enumerate(result[:4]):
                print(f'    Q{i+1}: {q[:80]}', flush=True)
            return result[:3]
    except Exception as e:
        print(f'  ⚠️ Query rewrite failed: {e}', flush=True)
    
    return [question]


def _multi_query_search(question, history=None, top_k=15, allowed_doc_ids=None):
    """Multi-query retrieval: rewrite query → search each → merge with RRF."""
    queries = _rewrite_query(question, history)
    
    if len(queries) <= 1:
        return search_docs(question, top_k=top_k, allowed_doc_ids=allowed_doc_ids)
    
    # Search with each query
    all_results = {}  # chunk_key -> {chunk_data, ranks: []}
    
    for qi, q in enumerate(queries):
        hits = search_docs(q, top_k=top_k, allowed_doc_ids=allowed_doc_ids)
        for rank, h in enumerate(hits):
            key = f'{h["docId"]}_{h["chunkIndex"]}'
            if key not in all_results:
                all_results[key] = {'data': h, 'ranks': []}
            all_results[key]['ranks'].append(rank)
    
    # RRF fusion across queries
    k_rrf = 60
    fused = []
    for key, item in all_results.items():
        rrf_score = sum(1.0 / (k_rrf + r) for r in item['ranks'])
        # Boost chunks found by multiple queries
        query_coverage = len(item['ranks']) / len(queries)
        rrf_score *= (1 + query_coverage * 0.5)
        item['data']['score'] = round(rrf_score, 6)
        item['data']['queryHits'] = len(item['ranks'])
        fused.append(item['data'])
    
    fused.sort(key=lambda x: x['score'], reverse=True)
    print(f'  🔀 Multi-query fusion: {len(fused)} unique chunks from {len(queries)} queries', flush=True)
    return fused[:top_k + 5]  # Return extra for better coverage


# ═══════════════════════════════════════════════════════
#  Level 4: AI Reranker
# ═══════════════════════════════════════════════════════

def _ai_rerank(question, chunks, top_k=15):
    """Use AI to rerank search results by relevance scoring."""
    if not chunks or not GEMINI_KEYS:
        return chunks[:top_k]
    
    # Only rerank top candidates (save tokens)
    candidates = chunks[:min(15, len(chunks))]
    
    # Include summary + chunk text for better relevance scoring
    def _rerank_text(c, max_len=600):
        s = c.get('summary', '')
        t = c['chunk']
        if s:
            return f'สรุป: {s}\nเนื้อหา: {t[:max_len - len(s)]}'
        return t[:max_len]
    
    chunk_texts = '\n\n'.join(
        f'[Doc {i+1}]\n{_rerank_text(c)}' 
        for i, c in enumerate(candidates)
    )
    
    prompt = f"""คุณเป็นผู้เชี่ยวชาญจัดอันดับความเกี่ยวข้องของเอกสาร

คำถาม: {question}

เอกสาร:
{chunk_texts}

ให้คะแนนแต่ละ Doc (1-10) ตามความเกี่ยวข้องกับคำถาม
- 10 = ตอบคำถามได้ตรงเป๊ะ
- 7-9 = เกี่ยวข้องมาก
- 4-6 = เกี่ยวข้องบ้าง
- 1-3 = ไม่ค่อยเกี่ยว

ตอบเป็น JSON array เท่านั้น: [{{"doc":1,"score":8}},{{"doc":2,"score":5}}]"""
    
    try:
        resp = _gemini_call(prompt, is_embed=False)
        resp_clean = resp.strip()
        if '```' in resp_clean:
            m = re.search(r'```(?:json)?\s*(.*?)```', resp_clean, re.DOTALL)
            if m:
                resp_clean = m.group(1)
        scores = json.loads(resp_clean)
        
        # Apply AI scores
        score_map = {}
        for s in scores:
            idx = s.get('doc', 0) - 1
            if 0 <= idx < len(candidates):
                score_map[idx] = s.get('score', 5)
        
        # Combine original score + AI rerank score
        for idx, c in enumerate(candidates):
            ai_score = score_map.get(idx, 5)
            original_score = c.get('score', 0)
            # Weighted combination: 60% AI rerank + 40% original
            c['rerank_score'] = round(ai_score / 10.0 * 0.6 + min(original_score * 50, 1.0) * 0.4, 4)
            c['ai_relevance'] = ai_score
        
        # Sort by rerank score
        candidates.sort(key=lambda x: x.get('rerank_score', 0), reverse=True)
        
        # Keep all chunks — let the AI answer model decide relevance
        # (aggressive filtering here was dropping important context)
        reranked = candidates
        
        print(f'  🏆 Reranked: {len(reranked)} chunks (filtered from {len(candidates)})', flush=True)
        for i, c in enumerate(reranked[:5]):
            print(f'    [{i+1}] ai={c.get("ai_relevance","?")} rerank={c.get("rerank_score",0):.3f} {c["chunk"][:50]}...', flush=True)
        
        return reranked[:top_k]
    
    except Exception as e:
        print(f'  ⚠️ Reranker failed: {e}', flush=True)
        return chunks[:top_k]


# ═══════════════════════════════════════════════════════
#  Level 4: Contextual Compression
# ═══════════════════════════════════════════════════════

def _compress_context(question, chunks):
    """Compress retrieved chunks — keep only relevant parts for the question."""
    if not chunks or not GEMINI_KEYS or len(chunks) <= 2:
        return chunks
    
    chunk_texts = '\n\n'.join(
        f'[Chunk {i+1}]\n{c["chunk"][:500]}' 
        for i, c in enumerate(chunks[:8])
    )
    
    prompt = f"""คุณเป็นผู้เชี่ยวชาญสกัดข้อมูลจากเอกสาร

คำถาม: {question}

เอกสาร:
{chunk_texts}

งาน: สกัดเฉพาะส่วนที่เกี่ยวข้องกับคำถามจากแต่ละ Chunk ตัดส่วนที่ไม่เกี่ยวออก
ถ้า Chunk ไม่เกี่ยวข้องเลย ให้ข้าม

ตอบเป็น JSON array: [{{"chunk":1,"compressed":"ข้อความที่สกัดแล้ว..."}}]"""
    
    try:
        resp = _gemini_call(prompt, is_embed=False)
        resp_clean = resp.strip()
        if '```' in resp_clean:
            m = re.search(r'```(?:json)?\s*(.*?)```', resp_clean, re.DOTALL)
            if m:
                resp_clean = m.group(1)
        compressed = json.loads(resp_clean)
        
        # Apply compressed text
        compress_map = {}
        for item in compressed:
            idx = item.get('chunk', 0) - 1
            if 0 <= idx < len(chunks) and item.get('compressed'):
                compress_map[idx] = item['compressed']
        
        result = []
        for idx, c in enumerate(chunks[:8]):
            if idx in compress_map:
                c_copy = dict(c)
                c_copy['original_chunk'] = c['chunk']
                c_copy['chunk'] = compress_map[idx]
                result.append(c_copy)
            elif idx < 3:  # Keep top 3 even if not compressed
                result.append(c)
        
        total_orig = sum(len(c['chunk']) for c in chunks[:8])
        total_comp = sum(len(c['chunk']) for c in result)
        ratio = total_comp / max(total_orig, 1) * 100
        print(f'  📦 Compressed: {len(result)} chunks, {ratio:.0f}% of original size', flush=True)
        
        return result if result else chunks[:8]
    
    except Exception as e:
        print(f'  ⚠️ Compression failed: {e}', flush=True)
        return chunks[:8]


# ═══════════════════════════════════════════════════════
#  Level 5: Agentic RAG — Router + Decompose + Self-Reflect
# ═══════════════════════════════════════════════════════

def _agent_classify(question, history=None):
    """Classify question type and decide retrieval strategy."""
    if not GEMINI_KEYS:
        return {'type': 'simple', 'strategy': 'search', 'sub_questions': []}

    hist_ctx = ''
    if history and len(history) > 0:
        hist_ctx = '\nบทสนทนาก่อนหน้า:\n' + '\n'.join(
            f'User: {h.get("question","")}' for h in history[-3:]
        )

    prompt = f"""คุณเป็น AI Router สำหรับระบบค้นหาเอกสาร
วิเคราะห์คำถามของผู้ใช้แล้วจำแนกประเภท
{hist_ctx}
คำถาม: {question}

จำแนกเป็น 1 ใน 4 ประเภท:
1. "simple" = คำถามตรงๆ ค้นครั้งเดียวตอบได้ (เช่น "วงเงินเท่าไหร่")
2. "multi_step" = คำถามซับซ้อน ต้องค้นหลายส่วนแล้วรวมคำตอบ (เช่น "เปรียบเทียบ A กับ B", "สรุปทั้งหมด")
3. "analytical" = ต้องวิเคราะห์/คำนวณ/สรุป (เช่น "มีกี่รายการ", "รวมเท่าไหร่")
4. "conversational" = ถามต่อจากบทสนทนาก่อนหน้า ไม่ต้องค้นใหม่มาก

ถ้าเป็น multi_step ให้แตกคำถามย่อย (sub_questions) ที่ต้องค้นหาแยกกัน (2-4 คำถาม)

ตอบ JSON เท่านั้น: {{"type":"simple|multi_step|analytical|conversational","strategy":"search","sub_questions":["q1","q2"],"reasoning":"เหตุผลสั้นๆ"}}"""

    try:
        resp = _gemini_call(prompt, is_embed=False)
        resp_clean = resp.strip()
        if '```' in resp_clean:
            m = re.search(r'```(?:json)?\s*(.*?)```', resp_clean, re.DOTALL)
            if m: resp_clean = m.group(1)
        result = json.loads(resp_clean)
        qtype = result.get('type', 'simple')
        print(f'  🧠 Agent classified: {qtype} — {result.get("reasoning","")}', flush=True)
        return result
    except Exception as e:
        print(f'  ⚠️ Agent classify failed: {e}', flush=True)
        return {'type': 'simple', 'strategy': 'search', 'sub_questions': []}


def _agent_decompose_search(question, sub_questions, history=None, top_k=10, allowed_doc_ids=None):
    """Multi-step: search each sub-question separately, then merge."""
    all_chunks = {}
    step_results = []

    for sq in sub_questions[:4]:
        print(f'    🔍 Sub-search: "{sq[:60]}"', flush=True)
        hits = search_docs(sq, top_k=top_k, allowed_doc_ids=allowed_doc_ids)
        step_results.append({'question': sq, 'hits': len(hits)})
        for rank, h in enumerate(hits):
            key = f'{h["docId"]}_{h["chunkIndex"]}'
            if key not in all_chunks:
                all_chunks[key] = {'data': h, 'ranks': [], 'from_questions': []}
            all_chunks[key]['ranks'].append(rank)
            all_chunks[key]['from_questions'].append(sq[:30])

    # RRF fusion across sub-questions
    k_rrf = 60
    fused = []
    for key, item in all_chunks.items():
        rrf_score = sum(1.0 / (k_rrf + r) for r in item['ranks'])
        coverage = len(set(item['from_questions'])) / max(len(sub_questions), 1)
        rrf_score *= (1 + coverage * 0.8)  # big boost for multi-question coverage
        item['data']['score'] = round(rrf_score, 6)
        item['data']['queryHits'] = len(item['ranks'])
        fused.append(item['data'])

    fused.sort(key=lambda x: x['score'], reverse=True)
    print(f'  🔀 Decomposed search: {len(fused)} chunks from {len(sub_questions)} sub-questions', flush=True)
    return fused[:top_k * 2], step_results


def _agent_reflect(question, answer, chunks):
    """Self-RAG: evaluate answer quality, decide if re-search needed."""
    if not GEMINI_KEYS or not answer:
        return {'quality': 'ok', 'score': 7, 'needs_more': False, 'feedback': ''}

    chunk_summary = ', '.join(set(c.get('docName','') for c in chunks[:5]))

    prompt = f"""คุณเป็น AI ผู้ตรวจสอบคุณภาพคำตอบ

คำถาม: {question}
คำตอบ: {answer[:1500]}
เอกสารที่ใช้: {chunk_summary}

ประเมินคำตอบ:
1. ตอบคำถามได้ครบถ้วนหรือไม่?
2. ข้อมูลถูกต้องตรงกับเอกสารหรือไม่?
3. มีส่วนที่ขาดหายหรือต้องค้นเพิ่มไหม?

ให้คะแนน 1-10 และบอกว่าต้องค้นเพิ่มไหม

ตอบ JSON เท่านั้น: {{"score":8,"quality":"good|ok|poor","needs_more":false,"missing":"","feedback":"สั้นๆ"}}"""

    try:
        resp = _gemini_call(prompt, is_embed=False)
        resp_clean = resp.strip()
        if '```' in resp_clean:
            m = re.search(r'```(?:json)?\s*(.*?)```', resp_clean, re.DOTALL)
            if m: resp_clean = m.group(1)
        result = json.loads(resp_clean)
        print(f'  🪞 Reflection: score={result.get("score","?")}/10 quality={result.get("quality","?")} needs_more={result.get("needs_more",False)}', flush=True)
        if result.get('feedback'):
            print(f'     Feedback: {result["feedback"][:100]}', flush=True)
        return result
    except Exception as e:
        print(f'  ⚠️ Reflection failed: {e}', flush=True)
        return {'quality': 'ok', 'score': 7, 'needs_more': False, 'feedback': ''}


def _agent_refine_answer(question, original_answer, extra_chunks, reflection):
    """Refine answer with additional context from re-search."""
    if not GEMINI_KEYS or not extra_chunks:
        return original_answer

    extra_ctx = '\n\n'.join(f'[เพิ่มเติม {i+1}] {c["chunk"][:400]}' for i, c in enumerate(extra_chunks[:5]))

    prompt = f"""คำถาม: {question}

คำตอบเดิม:
{original_answer[:2000]}

ข้อมูลเพิ่มเติม:
{extra_ctx}

ข้อเสนอแนะ: {reflection.get('feedback','')} / ส่วนที่ขาด: {reflection.get('missing','')}

ปรับปรุงคำตอบให้ครบถ้วนขึ้น โดยรวมข้อมูลเพิ่มเติมเข้าไป ตอบภาษาไทย จัดรูปแบบ Markdown
ห้ามอ้างอิงชื่อเอกสาร ห้ามบอกที่มา ตอบเนื้อหาอย่างเดียว"""

    try:
        refined = _gemini_call(prompt, is_embed=False)
        if refined and len(refined) > len(original_answer) * 0.5:
            print(f'  ✨ Answer refined: {len(original_answer)} → {len(refined)} chars', flush=True)
            return refined
    except Exception as e:
        print(f'  ⚠️ Refine failed: {e}', flush=True)

    return original_answer



def _level4_pipeline(question, history=None, allowed_doc_ids=None):
    """Document-First RAG: Search chunks to find docs → Send FULL doc text to AI.
    
    Why: Gemini 2.5 Flash has 1M token context. A typical document is 10-100K chars.
    Sending the full document eliminates ALL information loss from chunking.
    No reranker, no neighbor expansion, no compression hacks needed.
    """
    t_total = time.time()
    MAX_CONTEXT_CHARS = 500000  # ~150K tokens, well within Gemini's 1M

    # ── Fast-path: greetings / chitchat — skip RAG entirely ──
    if _is_chitchat(question):
        print(f'  💬 Chitchat fast-path: "{question}"', flush=True)
        return _chitchat_respond(question, history)

    # ── Step 1: Quick search to identify relevant DOCUMENTS ──
    t0 = time.time()
    hits = search_docs(question, top_k=20, allowed_doc_ids=allowed_doc_ids)
    t_search = time.time() - t0
    
    # Hybrid mode: no docs uploaded OR no hits at all → general knowledge
    if not documents:
        return _general_respond(question, history, reason='no_documents')
    if not hits:
        return _general_respond(question, history, reason='no_search_hits')

    # Note: RRF scores are very small (0.01-0.05). We do NOT pre-filter by score
    # because false-negatives are worse than running RAG once. Instead we detect
    # "no info in docs" from the AI answer itself (post-check below) and fallback.

    # ── Step 2: Identify top documents (not chunks) ──
    doc_scores = {}
    for h in hits:
        did = h['docId']
        if did not in doc_scores:
            doc_scores[did] = {'score': 0, 'name': h['docName'], 'hits': 0}
        doc_scores[did]['score'] += h.get('score', 0)
        doc_scores[did]['hits'] += 1
    
    # Rank documents by combined score
    ranked_docs = sorted(doc_scores.items(), key=lambda x: x[1]['score'], reverse=True)
    
    # ── Step 3: Load FULL TEXT for top documents ──
    context_parts = []
    total_chars = 0
    docs_used = []
    
    for doc_id, info in ranked_docs:
        # Try to load full text from cache
        full_text = _load_cached_text(doc_id)
        
        if full_text and total_chars + len(full_text) <= MAX_CONTEXT_CHARS:
            # Use FULL document text — no information loss
            doc_name = info['name']
            context_parts.append(f'══ เอกสาร: {doc_name} ══\n\n{full_text}')
            total_chars += len(full_text)
            docs_used.append(doc_name)
            print(f'  📄 Full doc: {doc_name} ({len(full_text):,} chars)', flush=True)
        elif full_text:
            # Document too large for remaining space — use chunk excerpts
            doc_chunks = _db_get_doc_chunks(doc_id) or enriched_store.get(doc_id, {}).get('chunks', [])
            if doc_chunks:
                # Get chunks relevant to this question, sorted by position
                relevant_idxs = set(h['chunkIndex'] for h in hits if h['docId'] == doc_id)
                # Also include surrounding chunks for context
                expanded = set()
                for idx in relevant_idxs:
                    for offset in range(-3, 4):
                        ni = idx + offset
                        if 0 <= ni < len(doc_chunks):
                            expanded.add(ni)
                
                sorted_idxs = sorted(expanded)
                excerpt_parts = []
                for ci in sorted_idxs:
                    text = doc_chunks[ci].get('text', '')
                    excerpt_parts.append(text)
                    total_chars += len(text)
                    if total_chars > MAX_CONTEXT_CHARS:
                        break
                
                doc_name = info['name']
                context_parts.append(f'══ เอกสาร: {doc_name} (ส่วนที่เกี่ยวข้อง) ══\n\n' + '\n\n'.join(excerpt_parts))
                docs_used.append(doc_name)
                print(f'  📄 Excerpt: {doc_name} ({len(sorted_idxs)} chunks)', flush=True)
        
        if total_chars >= MAX_CONTEXT_CHARS:
            break
    
    if not context_parts:
        return _local_respond(question, hits)
    
    # ── Step 4: Build prompt and get AI answer ──
    t0 = time.time()
    full_context = '\n\n'.join(context_parts)
    
    prompt = f'{SYSTEM_PROMPT}\n\nเอกสาร:\n\n{full_context}\n\n'
    
    if history:
        for h in history[-2:]:
            prompt += f'\nUser: {h.get("question", "")}\nAI: {h.get("answer", "")[:300]}\n'
    
    prompt += f'\nคำถาม: {question}\n\nตอบจากเอกสารให้ครบถ้วน จัดรูปแบบด้วย Markdown'
    
    try:
        answer = _gemini_call(prompt, is_embed=False, model_name=GEMINI_CHAT_MODEL)
        t_answer = time.time() - t0
        total_time = time.time() - t_total
        
        print(f'  ✅ Doc-First RAG: {len(docs_used)} docs, {total_chars:,} chars context | {total_time:.1f}s', flush=True)
        
        # Hybrid post-check: if AI says "not in document", fallback to general knowledge
        not_found_markers = [
            'ไม่พบข้อมูล', 'ไม่พบในเอกสาร', 'ไม่มีข้อมูล', 'ไม่ได้ระบุ',
            'ไม่มีรายละเอียด', 'ไม่ปรากฏ', 'เอกสารไม่ได้กล่าวถึง',
            'ไม่ได้กล่าวถึง', 'ไม่ได้พูดถึง', 'ไม่มีเนื้อหา',
            'no information', 'not mentioned', 'not found in', 'no data',
            'document does not', 'documents do not',
        ]
        ans_lower = (answer or '').lower()
        is_not_found = any(m.lower() in ans_lower for m in not_found_markers)
        # Only fallback if the answer is short AND clearly says "not found"
        if is_not_found and len(answer) < 600:
            print(f'  💭 AI says "not in docs" → general knowledge fallback', flush=True)
            general = _general_respond(question, history, reason='not_in_docs')
            general['answer'] = ('> 💭 **ไม่พบข้อมูลในเอกสาร — ตอบจากความรู้ทั่วไปแทน**\n\n'
                                 + general['answer'].split('\n\n', 1)[-1])
            return general
        return {
            'answer': answer,
            'model': GEMINI_CHAT_MODEL,
            'provider': 'Google',
            'sources': [{'docName': d, 'preview': '', 'score': 0} for d in docs_used],
            'searchMethod': 'document-first-rag',
            'pipelineTime': round(total_time, 1),
        }
    except Exception as e:
        print(f'  ❌ AI Error: {e}', flush=True)
        return _local_respond(question, hits, f'AI error: {e}')


def _is_chitchat(q):
    """Detect greetings / small-talk / meta questions that don't need RAG."""
    if not q: return False
    qs = q.strip().lower()
    if len(qs) > 60: return False
    # Thai + English greetings / chitchat patterns
    patterns = [
        'สวัสดี', 'หวัดดี', 'ดีจ้า', 'ดีครับ', 'ดีค่ะ',
        'ขอบคุณ', 'ขอบใจ', 'ทดสอบ', 'เทส', 'ลองดู',
        'คุณคือใคร', 'คุณเป็นใคร', 'แนะนำตัว', 'ทำอะไรได้',
        'ช่วยอะไรได้', 'ใช้งานยังไง', 'มีฟีเจอร์',
        'hello', 'hi ', 'hey', 'thanks', 'thank you', 'who are you',
        'what can you do', 'how do you work', 'test', 'testing',
    ]
    # Exact greeting (entire message is just "สวัสดี" / "hi")
    short_greets = {'สวัสดี','หวัดดี','hi','hello','hey','ดี','ทดสอบ','test'}
    if qs in short_greets: return True
    if qs.rstrip('?!.ครับค่ะนะจ้า ') in short_greets: return True
    # Starts-with patterns + very short
    if len(qs) <= 30:
        for pat in patterns:
            if qs.startswith(pat) or qs == pat.strip():
                return True
    return False


def _chitchat_respond(question, history=None):
    """Lightweight greeting responder — 1 Gemini call, no retrieval."""
    doc_count = len(documents)
    doc_names = ', '.join(d.get('name','') for d in documents[:5]) if documents else ''
    extra = f' (มี {doc_count} เอกสาร: {doc_names})' if doc_count else ' (ยังไม่มีเอกสาร — ให้ผู้ใช้ upload ก่อน)'
    prompt = f"""คุณคือผู้ช่วย AI วิเคราะห์เอกสาร{extra}
ผู้ใช้ทักทาย/ถามทั่วไปว่า: "{question}"
ตอบสั้นๆ เป็นกันเอง 1-2 ประโยค ไม่ต้องสรุปเอกสารใดๆ
- ถ้าเป็นการทักทาย → ทักทายกลับ + บอกว่าช่วยอะไรได้
- ถ้าถามว่าทำอะไรได้ → บอกสั้นๆ ว่าวิเคราะห์เอกสาร PDF/Word/Excel ตอบคำถามจากเนื้อหาได้
- ถ้าขอบคุณ → ตอบรับสั้นๆ
ห้าม dump เนื้อหาเอกสาร ห้ามใช้หัวข้อหรือ bullet มาก ตอบแบบสนทนาธรรมดา"""
    try:
        ans = _gemini_call(prompt, is_embed=False).strip()
    except Exception as e:
        ans = 'สวัสดีครับ! ผมช่วยวิเคราะห์เอกสารและตอบคำถามจากเนื้อหาได้ ลองอัปโหลดเอกสารแล้วถามได้เลยครับ'
    return {
        'answer': ans,
        'model': GEMINI_CHAT_MODEL,
        'provider': 'Google',
        'sources': [],
        'searchMethod': 'chitchat-fastpath',
        'agentType': 'chitchat',
        'qualityScore': 10,
        'agentLog': ['chitchat:fast'],
        'pipelineTime': 0.0,
        'documentCount': doc_count,
        'totalChunks': 0,
        'hitCount': 0,
    }


def _general_respond(question, history=None, reason='no_relevant_docs'):
    """Hybrid fallback: answer from Gemini's general knowledge when no doc context is found.
    The answer is clearly badged so the user knows it is NOT from their uploaded documents.
    """
    hist_ctx = ''
    if history:
        last = history[-3:]
        hist_ctx = '\n\nบทสนทนาก่อนหน้า:\n' + '\n'.join(
            f'User: {h.get("question","")}\nAI: {(h.get("answer","") or "")[:300]}'
            for h in last
        )
    doc_count = len(documents)
    prompt = f"""คุณคือผู้ช่วย AI ทั่วไป (Google Gemini)
ผู้ใช้ถามคำถามที่ไม่เกี่ยวข้องกับเอกสารที่อัปโหลดไว้ ({doc_count} ไฟล์) — จึงตอบจากความรู้ทั่วไปได้

คำถาม: {question}{hist_ctx}

ตอบเป็นภาษาเดียวกับคำถาม กระชับ ตรงประเด็น ใช้ markdown ได้
หากเป็นคำถามที่ต้องอ้างอิงข้อเท็จจริงล่าสุด ให้บอกว่าข้อมูลอาจไม่เป็นปัจจุบัน
ห้ามเริ่มต้นด้วย "จากเอกสาร" — เพราะคำตอบนี้ไม่ได้มาจากเอกสาร"""
    try:
        ans = _gemini_call(prompt, is_embed=False).strip()
    except Exception as e:
        ans = f'ขออภัย ไม่สามารถตอบคำถามนี้ได้ในตอนนี้ ({e})'
    badge = '> 💭 **ตอบจากความรู้ทั่วไป (ไม่ได้มาจากเอกสารที่อัปโหลด)**\n\n'
    return {
        'answer': badge + ans,
        'model': GEMINI_CHAT_MODEL,
        'provider': 'Google',
        'sources': [],
        'searchMethod': 'general-knowledge',
        'agentType': 'general',
        'fallbackReason': reason,
        'qualityScore': 7,
        'agentLog': [f'general:{reason}'],
        'pipelineTime': 0.0,
        'documentCount': doc_count,
        'totalChunks': 0,
        'hitCount': 0,
    }


def _agentic_pipeline(question, history=None, allowed_doc_ids=None):
    """Level 5 Agentic RAG: Route → Search → Rerank → Compress → Answer → Reflect → Refine."""
    t_total = time.time()
    agent_log = []

    # ── Fast-path: greetings / chitchat — skip RAG entirely ──
    if _is_chitchat(question):
        print(f'  💬 Chitchat fast-path: "{question}"', flush=True)
        return _chitchat_respond(question, history)

    # ── Step 1: Agent Classification ──
    t0 = time.time()
    classification = _agent_classify(question, history)
    qtype = classification.get('type', 'simple')
    sub_questions = classification.get('sub_questions', [])
    agent_log.append(f'classify:{qtype}:{time.time()-t0:.1f}s')

    # ── Step 2: Smart Retrieval based on type ──
    t0 = time.time()
    step_results = []
    if qtype == 'multi_step' and sub_questions:
        # Decomposed multi-step search
        hits, step_results = _agent_decompose_search(question, sub_questions, history, top_k=10, allowed_doc_ids=allowed_doc_ids)
        print(f'  📊 Multi-step found {len(hits)} chunks ({time.time()-t0:.1f}s)', flush=True)
    elif qtype == 'conversational' and history:
        # Light search + rely on history
        hits = search_docs(question, top_k=8, allowed_doc_ids=allowed_doc_ids)
        print(f'  📊 Conversational search: {len(hits)} chunks ({time.time()-t0:.1f}s)', flush=True)
    else:
        # Multi-query search (Level 4)
        hits = _multi_query_search(question, history, top_k=15, allowed_doc_ids=allowed_doc_ids)
        print(f'  📊 Multi-query found {len(hits)} chunks ({time.time()-t0:.1f}s)', flush=True)
    agent_log.append(f'search:{len(hits)}hits:{time.time()-t0:.1f}s')

    # ── Step 3: AI Reranker ──
    t0 = time.time()
    hits = _ai_rerank(question, hits, top_k=8)
    agent_log.append(f'rerank:{len(hits)}:{time.time()-t0:.1f}s')

    # ── Step 4: Contextual Compression ──
    t0 = time.time()
    compressed = _compress_context(question, hits)
    agent_log.append(f'compress:{len(compressed)}:{time.time()-t0:.1f}s')

    # ── Step 5: Generate Answer ──
    t0 = time.time()
    result = ai_respond(question, compressed, history)
    answer = result.get('answer', '')
    agent_log.append(f'answer:{len(answer)}chars:{time.time()-t0:.1f}s')

    # ── Step 6: Self-Reflection ──
    t0 = time.time()
    reflection = _agent_reflect(question, answer, compressed)
    quality_score = reflection.get('score', 7)
    agent_log.append(f'reflect:score={quality_score}:{time.time()-t0:.1f}s')

    # ── Step 7: Re-search + Refine if quality is low ──
    if reflection.get('needs_more') and quality_score <= 5:
        print(f'  🔄 Quality low ({quality_score}/10), re-searching...', flush=True)
        t0 = time.time()
        missing = reflection.get('missing', question)
        extra_query = missing if missing else question
        extra_hits = search_docs(extra_query, top_k=8, allowed_doc_ids=allowed_doc_ids)
        if extra_hits:
            extra_reranked = _ai_rerank(extra_query, extra_hits, top_k=5)
            refined_answer = _agent_refine_answer(question, answer, extra_reranked, reflection)
            result['answer'] = refined_answer
            result['refined'] = True
            agent_log.append(f'refine:{len(refined_answer)}chars:{time.time()-t0:.1f}s')
        else:
            agent_log.append(f'refine:no_extra_hits:{time.time()-t0:.1f}s')
    else:
        result['refined'] = False

    total_time = time.time() - t_total
    result['searchMethod'] = 'level5-agentic'
    result['agentType'] = qtype
    result['qualityScore'] = quality_score
    result['agentLog'] = agent_log
    result['pipelineTime'] = round(total_time, 1)

    print(f'  🤖 Agentic pipeline done: {qtype} | quality={quality_score}/10 | refined={result["refined"]} | {total_time:.1f}s', flush=True)
    print(f'     Log: {" → ".join(agent_log)}', flush=True)

    return result


# ═══════════════════════════════════════════════════════
#  AI Response Generation
# ═══════════════════════════════════════════════════════

SYSTEM_PROMPT = """คุณเป็น AI ผู้เชี่ยวชาญวิเคราะห์เอกสารสำหรับหน่วยงาน
กฎ:
1. ตอบภาษาไทย สุภาพ ชัดเจน ครบถ้วน
2. ตอบจากข้อมูลในเอกสารที่ให้มาเท่านั้น ห้ามแต่งเพิ่ม
3. ห้ามอ้างอิงชื่อไฟล์ ห้ามบอกที่มา ตอบเนื้อหาอย่างเดียว
4. จัดรูปแบบด้วย Markdown ใช้ emoji พอเหมาะ
5. ถ้ามีรายการ ต้องแจกแจงครบทุกข้อ ห้ามสรุปรวบ ห้ามตัดทอน
6. ถ้ามีตัวเลข เปอร์เซ็นต์ จำนวนเงิน ระบุให้ถูกต้องตามเอกสาร
7. ถ้าไม่พบข้อมูล ให้บอกตรงๆ ว่าไม่พบในเอกสาร"""


def ai_respond(question, chunks, history=None):
    """Try Gemini with smart key rotation."""
    if not GEMINI_KEYS:
        return _local_respond(question, chunks)

    try:
        ctx = _build_context(chunks)
        prompt = f'{SYSTEM_PROMPT}\n\nเอกสารที่เกี่ยวข้อง (Advanced RAG):\n\n{ctx}\n\n'

        if history:
            for h in history[-2:]:
                ans_preview = h.get("answer", "")[:300]
                prompt += f'\nUser: {h.get("question", "")}\nAI: {ans_preview}\n'

        prompt += f'\nคำถาม: {question}\n\nวิเคราะห์ข้อมูลจากเอกสารอย่างละเอียด ตอบให้ครบถ้วนทุกรายการ จัดรูปแบบด้วย Markdown'

        answer = _gemini_call(prompt, is_embed=False, model_name=GEMINI_CHAT_MODEL)
        
        return {
            'answer': answer,
            'model': GEMINI_CHAT_MODEL,
            'provider': 'Google',
            'sources': [{'docName': c['docName'], 'preview': c['chunk'][:150],
                         'score': c.get('score', 0)} for c in chunks],
            'searchMethod': 'hybrid',
        }
    except Exception as e:
        print(f'[AI Error] {e}', flush=True)
        return _local_respond(question, chunks, f'AI error: {e}')


def _build_context(chunks, max_total_chars=30000):
    """Build rich context from search results. Sort by document order for coherent reading."""
    if not chunks:
        return '(ไม่พบข้อมูลที่เกี่ยวข้อง)'
    
    # Sort by chunk index for natural reading order (crucial for multi-part answers)
    sorted_chunks = sorted(chunks, key=lambda x: (x.get('docId',''), x.get('chunkIndex', 0)))
    
    parts = []
    total = 0
    for i, c in enumerate(sorted_chunks):
        summary = c.get('summary', '')
        chunk_text = c["chunk"][:1500]
        if summary:
            part = f'[ส่วน {c.get("chunkIndex",i)+1}] สรุป: {summary}\nเนื้อหา: {chunk_text}'
        else:
            part = f'[ส่วน {c.get("chunkIndex",i)+1}] {chunk_text}'
        total += len(part)
        if total > max_total_chars:
            break
        parts.append(part)
    return '\n\n'.join(parts)


def _local_respond(question, chunks, error=None):
    """Fallback: local search-based response."""
    if not chunks:
        return {
            'answer': '🔍 ไม่พบข้อมูลที่เกี่ยวข้องกับคำถามนี้ในเอกสาร\n\n💡 **ลองทำสิ่งนี้:**\n- ใช้คำค้นอื่น\n- อัปโหลดเอกสารเพิ่ม\n- ถามเป็นคำสำคัญ เช่น ชื่อโครงการ วงเงิน',
            'model': 'Local Search', 'provider': 'Local',
            'sources': [], 'searchMethod': 'none',
        }
    
    unique_docs = list({c['docName'] for c in chunks})
    ans = f'## 📊 ผลการค้นหา: "{question}"\n\n'
    ans += f'*ค้นพบจาก {len(unique_docs)} เอกสาร · {len(chunks)} ส่วนที่เกี่ยวข้อง*\n\n'
    
    for i, c in enumerate(chunks[:5]):
        ans += f'### 📄 {c["docName"]} (ส่วนที่ {c.get("chunkIndex",0)+1})\n'
        ans += f'> {c["chunk"][:500]}{"..." if len(c["chunk"]) > 500 else ""}\n\n'
    
    if error:
        ans += f'\n---\n⚠️ AI: {error} — ใช้ Local Search'
    
    return {
        'answer': ans,
        'model': 'Local Search', 'provider': 'Local',
        'sources': [{'docName': c['docName'], 'preview': c['chunk'][:150],
                     'score': c.get('score', 0)} for c in chunks],
        'searchMethod': 'fulltext',
    }


# ═══════════════════════════════════════════════════════
#  API Routes
# ═══════════════════════════════════════════════════════

@app.route('/')
def serve_index():
    return send_from_directory(BASE_DIR, 'index.html')

@app.route('/favicon.ico')
def favicon():
    return '', 204

@app.route('/<path:path>')
def serve_static(path):
    if path.startswith('api/'):
        return jsonify({'error': 'Not found'}), 404
    return send_from_directory(BASE_DIR, path)


# ─── Upload ───────────────────────────────────────────

@app.route('/api/upload', methods=['POST'])
def api_upload():
    if 'file' not in request.files:
        return jsonify({'error': 'ไม่พบไฟล์'}), 400
    f = request.files['file']
    if not f.filename:
        return jsonify({'error': 'ไม่ได้เลือกไฟล์'}), 400

    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in ALLOWED_EXT:
        return jsonify({'error': f'ไม่รองรับไฟล์ {ext}'}), 400

    fid = uuid.uuid4().hex[:8]
    saved = os.path.join(UPLOAD_FOLDER, f'{fid}{ext}')
    f.save(saved)

    size = os.path.getsize(saved)
    if size > MAX_FILE_SIZE:
        os.remove(saved)
        return jsonify({'error': 'ไฟล์ใหญ่เกิน 50 MB'}), 400

    # Register document immediately (no text yet)
    doc = {
        'id': fid, 'name': f.filename, 'ext': ext,
        'size': size, 'path': saved,
        'uploadedAt': datetime.now().strftime('%Y-%m-%d %H:%M'),
        'text': '', 'chunks': [],
        'wordCount': 0,
        'owner': getattr(g, 'sid', 'public'),
    }
    documents.append(doc)
    _save_metadata()
    _update_progress(fid, 'processing', 2, 'Starting pipeline...', 'upload', 'active', 'Saving file...')

    # Run EVERYTHING in background (extract + chunk + enrich + embed)
    def _full_pipeline(doc_id, filepath, ext, filename):
        try:
            t_start = time.time()
            
            # Initialize all steps
            _update_progress(doc_id, 'processing', 2, 'Starting pipeline...',
                           'upload', 'done', f'File saved ({size/(1024*1024):.1f} MB)')
            for sn, sl in [('extract', 'Extract Text'), ('chunk', 'Chunking'), 
                           ('enrich', 'AI Enrichment'), ('embed', 'Embedding'), ('wiki', 'Building Wiki'), ('save', 'Saving')]:
                _update_progress(doc_id, 'processing', 2, 'Starting...', sn, 'pending', '')

            # ── Step 1: Extract Text ──
            _update_progress(doc_id, 'processing', 5, 'Extracting text...', 'extract', 'active', 'Reading document...')
            text = extract_text(filepath, ext, doc_id=doc_id)
            word_count = len(text.split()) if text else 0
            t_extract = time.time() - t_start
            _update_progress(doc_id, 'processing', 80, f'{word_count:,} words extracted',
                           'extract', 'done', f'{word_count:,} words in {t_extract:.1f}s')

            # Update document
            for d in documents:
                if d['id'] == doc_id:
                    d['text'] = text
                    d['chunks'] = chunk_text(text)
                    d['wordCount'] = word_count
                    break
            _save_metadata()

            # ── Step 2: Chunking ──
            t_chunk = time.time()
            _update_progress(doc_id, 'processing', 82, 'Creating chunks...', 'chunk', 'active', 'Splitting text...')
            chunks_raw = chunk_text(text)
            chunks = [{'text': c, 'section': '', 'summary': '', 'questions': [], 'embedding': []} for c in chunks_raw]
            total = len(chunks)
            t_chunk_dur = time.time() - t_chunk
            _update_progress(doc_id, 'processing', 83, f'{total} chunks created',
                           'chunk', 'done', f'{total} chunks in {t_chunk_dur:.1f}s')

            # ── Step 3+4: Enrich + Embed (CONCURRENT — embed raw first, then patch enriched) ──
            t_enrich = time.time()
            enrich_done = [False]
            n_sum_result = [0]

            def _do_enrich():
                if SKIP_ENRICHMENT:
                    _update_progress(doc_id, 'processing', 84, 'Skipped (fast mode)',
                                     'enrich', 'done', 'Fast mode — using raw text')
                    enrich_done[0] = True
                    return
                if total <= 500 and GEMINI_KEYS:
                    # Wait for keys to recover after OCR/extract phase
                    busy = sum(1 for k in GEMINI_KEYS if key_usage.get(k, {}).get('cooldown_until', 0) > time.time())
                    if busy >= len(GEMINI_KEYS) - 1:
                        wait_time = 10
                        print(f'  ⏳ {busy}/{len(GEMINI_KEYS)} keys on cooldown, waiting {wait_time}s for recovery...', flush=True)
                        _update_progress(doc_id, 'processing', 84, f'Waiting for API keys to recover...',
                                         'enrich', 'active', f'{busy} keys cooling down...')
                        time.sleep(wait_time)
                    _update_progress(doc_id, 'processing', 84, 'AI enriching summaries...',
                                     'enrich', 'active', f'0/{total} chunks...')
                    result = _enrich_chunks_parallel(chunks, filename, doc_id=doc_id)
                    # Apply enrichment to chunks in-place
                    for i, c in enumerate(result):
                        if i < len(chunks):
                            chunks[i]['summary'] = c.get('summary', '')
                            chunks[i]['questions'] = c.get('questions', [])
                    n_sum_result[0] = sum(1 for c in chunks if c.get('summary'))
                    t_dur = time.time() - t_enrich
                    _update_progress(doc_id, 'processing', 85, f'{n_sum_result[0]}/{total} enriched',
                                     'enrich', 'done', f'{n_sum_result[0]}/{total} summaries in {t_dur:.1f}s')
                else:
                    _update_progress(doc_id, 'processing', 84, 'Skipped enrichment',
                                     'enrich', 'done', f'Skipped ({total} chunks)')
                enrich_done[0] = True

            def _do_embed():
                _update_progress(doc_id, 'processing', 85, 'Creating embeddings...',
                                 'embed', 'active', f'0/{total} vectors...')
                embed_texts = [c['text'][:2000] for c in chunks]
                embs = _batch_embed_gemini(embed_texts, doc_id=doc_id, total_chunks=total)
                for j, emb in enumerate(embs):
                    if j < len(chunks):
                        chunks[j]['embedding'] = emb

            # Run both in parallel
            t_enrich_thread = threading.Thread(target=_do_enrich, daemon=True)
            t_embed_thread = threading.Thread(target=_do_embed, daemon=True)
            t_enrich_thread.start()
            t_embed_thread.start()
            t_enrich_thread.join()
            t_embed_thread.join()

            # After both done: re-embed ONLY enriched chunks (the ones with summary)
            n_sum = n_sum_result[0]
            if n_sum > 0:
                enriched_idxs = [i for i, c in enumerate(chunks) if c.get('summary')]
                enriched_texts = [(chunks[i].get('summary','') + '\n' + chunks[i]['text'])[:2000] for i in enriched_idxs]
                if enriched_texts:
                    _update_progress(doc_id, 'processing', 93, f'Re-embedding {len(enriched_texts)} enriched...',
                                     'embed', 'active', f'Patching {len(enriched_texts)}/{total} vectors')
                    re_embs = _batch_embed_gemini(enriched_texts)
                    for k, idx in enumerate(enriched_idxs):
                        if k < len(re_embs):
                            chunks[idx]['embedding'] = re_embs[k]

            n_emb = sum(1 for c in chunks if c.get('embedding') and len(c['embedding']) > 0 and any(v != 0 for v in c['embedding'][:5]))
            t_embed_dur = time.time() - t_enrich
            _update_progress(doc_id, 'processing', 97, f'{n_emb} embeddings',
                             'embed', 'done', f'{n_emb}/{total} vectors in {t_embed_dur:.1f}s')

            # ── Step 5: Save ──
            _update_progress(doc_id, 'processing', 98, 'Saving...', 'save', 'active', 'Writing cache...')
            n_sum = sum(1 for c in chunks if c.get('summary'))
            doc_type = {'type': 'general'}
            _save_enriched_cache(doc_id, chunks, doc_type)
            enriched_store[doc_id] = {'docType': doc_type, 'chunks': chunks}
            # Persist to SQLite + FAISS
            _db_upsert_chunks(doc_id, chunks)
            _faiss_remove(doc_id)
            _faiss_add(doc_id, [c.get('embedding', []) for c in chunks])
            for d in documents:
                if d['id'] == doc_id:
                    d['chunkCount'] = total
                    break
            _save_metadata()
            
            # ── Step 6: Build Wiki ──
            try:
                _update_progress(doc_id, 'processing', 99, 'Building wiki...', 'wiki', 'active', 'AI summarizing document...')
                t_wiki = time.time()
                _build_doc_wiki(doc_id, filename, chunks)
                _rebuild_wiki_index()
                _update_progress(doc_id, 'processing', 100, 'Wiki built', 'wiki', 'done', f'in {time.time()-t_wiki:.1f}s')
            except Exception as _we:
                print(f'  ⚠️ wiki step failed: {_we}', flush=True)
                _update_progress(doc_id, 'processing', 100, 'Wiki skipped', 'wiki', 'done', 'failed (non-fatal)')
            
            total_time = time.time() - t_start
            _update_progress(doc_id, 'done', 100, 'Ready!', 'save', 'done', f'Total: {total_time:.0f}s')
            t_enrich_total = time.time() - t_enrich - t_embed_dur if 't_embed_dur' in dir() else 0
            upload_progress[doc_id]['timing'] = {
                'total': round(total_time, 1),
                'extract': round(t_extract, 1),
                'enrich': round(t_enrich_total, 1),
                'embed': round(t_embed_dur, 1),
            }
            
            print(f'  ✅ Pipeline done: {filename} — {total} chunks, {n_sum} summaries, {n_emb} embeddings — ⏱️ {total_time:.1f}s', flush=True)

        except Exception as e:
            print(f'  ❌ Full pipeline error: {e}', flush=True)
            traceback.print_exc()
            _update_progress(doc_id, 'error', 0, str(e))

    thread = threading.Thread(target=_full_pipeline, args=(fid, saved, ext, f.filename), daemon=True)
    thread.start()

    return jsonify({
        'success': True,
        'document': {
            'id': doc['id'], 'name': doc['name'], 'ext': doc['ext'],
            'size': doc['size'], 'uploadedAt': doc['uploadedAt'],
            'wordCount': 0, 'chunkCount': 0,
        },
        'processing': True,
    })


# ─── Documents ────────────────────────────────────────

@app.route('/api/documents')
def api_documents():
    vis = _visible_docs()
    return jsonify({
        'documents': [{
            'id': d['id'], 'name': d['name'], 'ext': d['ext'],
            'size': d['size'], 'uploadedAt': d['uploadedAt'],
            'wordCount': d.get('wordCount', 0),
            'chunkCount': len(d.get('chunks', [])),
            'isPublic': (d.get('owner','public') == 'public'),
            'isMine': (d.get('owner') == getattr(g,'sid',None)),
        } for d in vis],
        'totalChunks': sum(len(d.get('chunks', [])) for d in vis),
        'totalWords': sum(d.get('wordCount', 0) for d in vis),
        'sessionId': getattr(g,'sid','')[:8],
    })

@app.route('/api/documents/<doc_id>', methods=['DELETE'])
def api_delete_doc(doc_id):
    global documents
    doc = next((d for d in documents if d['id'] == doc_id), None)
    if not doc:
        return jsonify({'error': 'ไม่พบเอกสาร'}), 404
    if doc.get('owner','public') != getattr(g,'sid',None):
        return jsonify({'error': 'ไม่มีสิทธิ์ลบเอกสารนี้ (เป็นของผู้อื่นหรือเอกสารสาธารณะ)'}), 403
    try:
        os.remove(doc['path'])
    except OSError:
        pass
    # Remove caches
    for suffix in ['.txt', '_enriched.json']:
        pth = os.path.join(TEXT_CACHE, f'{doc_id}{suffix}')
        if os.path.exists(pth):
            try: os.remove(pth)
            except: pass
    enriched_store.pop(doc_id, None)
    _db_remove_doc(doc_id)
    _faiss_remove(doc_id)
    documents = [d for d in documents if d['id'] != doc_id]
    _save_metadata()
    # Wiki cleanup (non-fatal)
    try:
        _delete_doc_wiki(doc_id)
        _rebuild_wiki_index()
    except Exception as _e:
        print(f'  WARN wiki cleanup on delete failed: {_e}', flush=True)
    return jsonify({'success': True})

@app.route('/api/documents/<doc_id>/preview')
def api_preview(doc_id):
    doc = next((d for d in documents if d['id'] == doc_id), None)
    if not doc:
        return jsonify({'error': 'ไม่พบเอกสาร'}), 404
    return jsonify({
        'name': doc['name'],
        'text': doc.get('text', '')[:10000],
        'wordCount': doc.get('wordCount', 0),
        'chunkCount': len(doc.get('chunks', [])),
    })


# ─── Chat ─────────────────────────────────────────────

@app.route('/api/chat', methods=['POST'])
def api_chat():
    data = request.get_json(silent=True) or {}
    q = (data.get('question') or '').strip()
    if not q:
        return jsonify({'error': 'ไม่พบคำถาม'}), 400

    vis = _visible_docs()
    if not vis:
        return jsonify({
            'answer': '📂 **ยังไม่มีเอกสารของคุณในระบบ**\n\nกรุณาอัปโหลดเอกสาร (PDF, Word, TXT, Excel) ก่อน\nไฟล์ของคุณจะถูกแยกออกจากผู้ใช้คนอื่น 🔒',
            'model': 'System', 'provider': 'System',
            'sources': [], 'documentCount': 0,
        })

    # ═══ Level 4 RAG Pipeline (fast) ═══
    allowed = set(d['id'] for d in vis)
    print(f'\n  🔍 Level 4 RAG: "{q}" (sid={getattr(g,"sid","?")[:8]} docs={len(allowed)})', flush=True)

    result = _level4_pipeline(q, chat_history[-6:] if chat_history else None, allowed_doc_ids=allowed)

    chat_history.append({'question': q, 'answer': result.get('answer', ''), 'time': datetime.now().isoformat()})
    if len(chat_history) > 50:
        chat_history.pop(0)

    result['documentCount'] = len(vis)
    result['totalChunks'] = sum(len(d.get('chunks', [])) for d in vis)
    result['hitCount'] = len(result.get('sources', []))
    return jsonify(result)


# ─── Upload Progress ──────────────────────────────────

@app.route('/api/upload/progress/<doc_id>')
def api_upload_progress(doc_id):
    prog = upload_progress.get(doc_id, {'status': 'unknown', 'progress': 0, 'message': '', 'steps': []})
    result = {
        'status': prog.get('status', 'unknown'),
        'progress': prog.get('progress', 0),
        'message': prog.get('message', ''),
        'steps': [],
        'timing': prog.get('timing', {}),
    }
    # Build steps info (remove internal fields)
    for s in prog.get('steps', []):
        step = {'name': s['name'], 'status': s['status'], 'detail': s.get('detail', '')}
        if 'duration' in s:
            step['duration'] = s['duration']
        # Calculate elapsed for active steps
        if s['status'] == 'active' and 'startedAt' in s:
            step['elapsed'] = round(time.time() - s['startedAt'], 1)
        result['steps'].append(step)
    
    # Also return latest doc info
    doc = next((d for d in documents if d['id'] == doc_id), None)
    if doc:
        result['wordCount'] = doc.get('wordCount', 0)
        result['chunkCount'] = len(doc.get('chunks', []))
    return jsonify(result)


# ─── Chat History ─────────────────────────────────────

@app.route('/api/chat/history')
def api_chat_history():
    return jsonify({'history': chat_history[-20:]})

@app.route('/api/chat/clear', methods=['POST'])
def api_chat_clear():
    chat_history.clear()
    return jsonify({'success': True})


# ─── Re-index / Re-enrich ────────────────────────────

@app.route('/api/reindex', methods=['POST'])
def api_reindex():
    """Re-run enrichment pipeline for all documents."""
    results = []
    for doc in documents:
        fid = doc['id']
        # Clear enriched cache
        ep = os.path.join(TEXT_CACHE, f'{fid}_enriched.json')
        if os.path.exists(ep):
            os.remove(ep)
        enriched_store.pop(fid, None)
        _db_remove_doc(fid)
        _faiss_remove(fid)
        
        text = doc.get('text', '') or _load_cached_text(fid) or ''
        thread = threading.Thread(target=_run_pipeline, args=(fid, doc['name'], text), daemon=True)
        thread.start()
        results.append({'name': doc['name'], 'status': 'pipeline_started'})
    
    return jsonify({'success': True, 'results': results})


@app.route('/api/reprocess/<doc_id>', methods=['POST'])
def api_reprocess(doc_id):
    """Full re-process: clear ALL cache (text + enriched) and re-run from OCR/extraction."""
    doc = next((d for d in documents if d['id'] == doc_id), None)
    if not doc:
        return jsonify({'error': 'Document not found'}), 404
    
    filepath = doc.get('path', '')
    ext = doc.get('ext', '')
    filename = doc.get('name', '')
    
    if not filepath or not os.path.exists(filepath):
        return jsonify({'error': 'File not found on disk'}), 404
    
    # Backup old caches (restore if reprocess fails)
    text_cache = _get_text_cache_path(doc_id)
    text_backup = text_cache + '.bak'
    enriched_cache = os.path.join(TEXT_CACHE, f'{doc_id}_enriched.json')
    enriched_backup = enriched_cache + '.bak'
    
    if os.path.exists(text_cache):
        import shutil
        shutil.copy2(text_cache, text_backup)
        os.remove(text_cache)
        print(f'  🗑️ Cleared text cache for {filename} (backup saved)', flush=True)
    
    if os.path.exists(enriched_cache):
        import shutil
        shutil.copy2(enriched_cache, enriched_backup)
        os.remove(enriched_cache)
        print(f'  🗑️ Cleared enriched cache for {filename} (backup saved)', flush=True)
    
    old_enriched = enriched_store.pop(doc_id, None)
    _db_remove_doc(doc_id)
    _faiss_remove(doc_id)
    
    # Reset document stats
    doc['wordCount'] = 0
    doc['chunkCount'] = 0
    doc['chunks'] = []
    doc['text'] = ''
    _save_metadata()
    
    # Build progress-aware full pipeline
    size = doc.get('size', 0)
    _update_progress(doc_id, 'processing', 1, 'Re-processing from scratch...',
                    'upload', 'done', f'File: {size/(1024*1024):.1f} MB')
    
    def _reprocess_pipeline():
        try:
            t_start = time.time()
            for sn in ['extract', 'chunk', 'enrich', 'embed', 'save']:
                _update_progress(doc_id, 'processing', 1, 'Starting...', sn, 'pending', '')
            
            _update_progress(doc_id, 'processing', 3, 'Extracting text (OCR)...', 'extract', 'active', 'Reading document...')
            text = extract_text(filepath, ext, doc_id=doc_id)
            word_count = len(text.split()) if text else 0
            t_extract = time.time() - t_start
            _update_progress(doc_id, 'processing', 80, f'{word_count:,} words extracted',
                           'extract', 'done', f'{word_count:,} words in {t_extract:.1f}s')
            
            doc['text'] = text
            doc['wordCount'] = word_count
            _save_metadata()
            
            t_chunk = time.time()
            _update_progress(doc_id, 'processing', 82, 'Creating chunks...', 'chunk', 'active', 'Splitting text...')
            chunks_raw = chunk_text(text)
            chunks = [{'text': c, 'section': '', 'summary': '', 'questions': [], 'embedding': []} for c in chunks_raw]
            total = len(chunks)
            t_chunk_dur = time.time() - t_chunk
            _update_progress(doc_id, 'processing', 83, f'{total} chunks created',
                           'chunk', 'done', f'{total} chunks in {t_chunk_dur:.1f}s')
            
            # Enrich + Embed (CONCURRENT)
            t_enrich = time.time()
            enrich_done = [False]
            n_sum_result = [0]

            def _do_enrich_r():
                if SKIP_ENRICHMENT:
                    _update_progress(doc_id, 'processing', 84, 'Skipped (fast mode)',
                                   'enrich', 'done', 'Fast mode — using raw text')
                    enrich_done[0] = True
                    return
                if total <= 500 and GEMINI_KEYS:
                    busy = sum(1 for k in GEMINI_KEYS if key_usage.get(k, {}).get('cooldown_until', 0) > time.time())
                    if busy >= len(GEMINI_KEYS) - 1:
                        print(f'  ⏳ {busy}/{len(GEMINI_KEYS)} keys cooling, waiting 10s...', flush=True)
                        time.sleep(10)
                    _update_progress(doc_id, 'processing', 84, 'AI enriching...',
                                   'enrich', 'active', f'0/{total} chunks...')
                    result = _enrich_chunks_parallel(chunks, filename, doc_id=doc_id)
                    for i, c in enumerate(result):
                        if i < len(chunks):
                            chunks[i]['summary'] = c.get('summary', '')
                            chunks[i]['questions'] = c.get('questions', [])
                    n_sum_result[0] = sum(1 for c in chunks if c.get('summary'))
                    t_dur = time.time() - t_enrich
                    _update_progress(doc_id, 'processing', 85, f'{n_sum_result[0]}/{total} enriched',
                                   'enrich', 'done', f'{n_sum_result[0]}/{total} in {t_dur:.1f}s')
                else:
                    _update_progress(doc_id, 'processing', 84, 'Skipped',
                                   'enrich', 'done', f'Skipped ({total} chunks)')
                enrich_done[0] = True

            def _do_embed_r():
                _update_progress(doc_id, 'processing', 85, 'Embedding...',
                               'embed', 'active', f'0/{total} vectors...')
                embs = _batch_embed_gemini([c['text'][:2000] for c in chunks], doc_id=doc_id, total_chunks=total)
                for j, emb in enumerate(embs):
                    if j < len(chunks):
                        chunks[j]['embedding'] = emb

            t1 = threading.Thread(target=_do_enrich_r, daemon=True)
            t2 = threading.Thread(target=_do_embed_r, daemon=True)
            t1.start(); t2.start()
            t1.join(); t2.join()

            n_sum = n_sum_result[0]
            if n_sum > 0:
                eidxs = [i for i, c in enumerate(chunks) if c.get('summary')]
                etexts = [(chunks[i].get('summary','')+'\n'+chunks[i]['text'])[:2000] for i in eidxs]
                if etexts:
                    _update_progress(doc_id, 'processing', 93, f'Re-embedding {len(etexts)} enriched...',
                                   'embed', 'active', f'Patching {len(etexts)}/{total}')
                    re_embs = _batch_embed_gemini(etexts)
                    for k, idx in enumerate(eidxs):
                        if k < len(re_embs): chunks[idx]['embedding'] = re_embs[k]

            n_emb = sum(1 for c in chunks if c.get('embedding') and len(c['embedding']) > 0 and any(v != 0 for v in c['embedding'][:5]))
            t_embed_dur = time.time() - t_enrich
            _update_progress(doc_id, 'processing', 97, f'{n_emb} embeddings',
                           'embed', 'done', f'{n_emb}/{total} vectors in {t_embed_dur:.1f}s')
            
            # Save
            _update_progress(doc_id, 'processing', 98, 'Saving...', 'save', 'active', 'Writing cache...')
            doc_type = {'type': 'general'}
            _save_enriched_cache(doc_id, chunks, doc_type)
            enriched_store[doc_id] = {'docType': doc_type, 'chunks': chunks}
            # Persist to SQLite + FAISS
            _db_upsert_chunks(doc_id, chunks)
            _faiss_remove(doc_id)
            _faiss_add(doc_id, [c.get('embedding', []) for c in chunks])
            doc['chunkCount'] = total
            doc['chunks'] = chunk_text(text)
            _save_metadata()
            
            # Build wiki (non-fatal)
            try:
                _build_doc_wiki(doc_id, filename, chunks)
                _rebuild_wiki_index()
            except Exception as _we:
                print(f'  ⚠️ wiki rebuild after reprocess failed: {_we}', flush=True)
            
            total_time = time.time() - t_start
            _update_progress(doc_id, 'done', 100, 'Ready!', 'save', 'done', f'Total: {total_time:.0f}s')
            upload_progress[doc_id]['timing'] = {
                'total': round(total_time, 1),
                'extract': round(t_extract, 1),
                'enrich': round(time.time() - t_enrich, 1),
                'embed': round(t_embed_dur, 1),
            }
            # Clean up backup files on success
            for _bak in [text_backup, enriched_backup]:
                if os.path.exists(_bak):
                    os.remove(_bak)
            print(f'  ✅ Reprocess done: {filename} — {total} chunks, {n_sum} summaries, {n_emb} embeddings — ⏱️ {total_time:.1f}s', flush=True)
        except Exception as e:
            print(f'  ❌ Reprocess error: {e}', flush=True)
            # Restore backed-up cache on failure
            import shutil
            if os.path.exists(text_backup) and not os.path.exists(text_cache):
                shutil.move(text_backup, text_cache)
                print(f'  ♻️ Restored text cache backup for {filename}', flush=True)
            if os.path.exists(enriched_backup) and not os.path.exists(enriched_cache):
                shutil.move(enriched_backup, enriched_cache)
                # Reload enriched store from backup
                try:
                    with open(enriched_cache, 'r') as _f:
                        _data = json.load(_f)
                    enriched_store[doc_id] = _data
                    _db_upsert_chunks(doc_id, _data.get('chunks', []))
                    _faiss_remove(doc_id)
                    _faiss_add(doc_id, [c.get('embedding', []) for c in _data.get('chunks', [])])
                    print(f'  ♻️ Restored enriched cache backup for {filename}', flush=True)
                except:
                    pass
            import traceback
            traceback.print_exc()
            _update_progress(doc_id, 'error', 0, str(e))
    
    thread = threading.Thread(target=_reprocess_pipeline, daemon=True)
    thread.start()
    
    return jsonify({'success': True, 'docId': doc_id, 'name': filename, 'processing': True})


# ─── Config (API key management for end-users) ───────

ENV_PATH = os.path.join(BASE_DIR, '.env')

def _mask_key(k):
    if not k: return ''
    if len(k) <= 12: return '***'
    return k[:6] + '...' + k[-4:]

def _read_env_file():
    """Read .env into dict (preserves keys we manage)."""
    data = {}
    if not os.path.exists(ENV_PATH):
        return data
    try:
        with open(ENV_PATH, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith('#') or '=' not in line:
                    continue
                k, v = line.split('=', 1)
                data[k.strip()] = v.strip().strip('"').strip("'")
    except Exception as e:
        print(f'[config] read .env failed: {e}', flush=True)
    return data

def _write_env_file(updates):
    """Merge `updates` dict into .env, preserving other vars + comments."""
    lines = []
    seen = set()
    if os.path.exists(ENV_PATH):
        with open(ENV_PATH, 'r', encoding='utf-8') as f:
            for line in f:
                stripped = line.strip()
                if not stripped or stripped.startswith('#') or '=' not in stripped:
                    lines.append(line.rstrip('\n'))
                    continue
                k = stripped.split('=', 1)[0].strip()
                if k in updates:
                    lines.append(f'{k}={updates[k]}')
                    seen.add(k)
                else:
                    lines.append(line.rstrip('\n'))
    for k, v in updates.items():
        if k not in seen:
            lines.append(f'{k}={v}')
    with open(ENV_PATH, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines) + '\n')
    try:
        os.chmod(ENV_PATH, 0o600)
    except Exception:
        pass

def _reload_gemini_keys():
    """Re-read GEMINI_API_KEY[_2.._5] from os.environ and rebuild GEMINI_KEYS in-place."""
    new_keys = []
    for k in ['GEMINI_API_KEY', 'GEMINI_API_KEY_2', 'GEMINI_API_KEY_3', 'GEMINI_API_KEY_4', 'GEMINI_API_KEY_5']:
        v = os.environ.get(k, '').strip()
        if v:
            new_keys.append(v)
    GEMINI_KEYS.clear()
    GEMINI_KEYS.extend(new_keys)
    return len(GEMINI_KEYS)

@app.route('/api/config', methods=['GET'])
def api_config_get():
    """Return masked key status for the setup UI."""
    env = _read_env_file()
    keys = []
    for slot in ['GEMINI_API_KEY', 'GEMINI_API_KEY_2', 'GEMINI_API_KEY_3', 'GEMINI_API_KEY_4', 'GEMINI_API_KEY_5']:
        v = env.get(slot, '') or os.environ.get(slot, '')
        keys.append({'slot': slot, 'set': bool(v), 'masked': _mask_key(v)})
    return jsonify({
        'configured': bool(GEMINI_KEYS),
        'keyCount': len(GEMINI_KEYS),
        'slots': keys,
        'envPath': ENV_PATH,
    })

@app.route('/api/config/keys', methods=['POST'])
def api_config_set_keys():
    """Save Gemini keys to .env and hot-reload runtime.
    Body: {"GEMINI_API_KEY": "...", "GEMINI_API_KEY_2": "...", ...} (any subset).
    Empty string clears that slot.
    """
    body = request.get_json(silent=True) or {}
    allowed = {'GEMINI_API_KEY', 'GEMINI_API_KEY_2', 'GEMINI_API_KEY_3', 'GEMINI_API_KEY_4', 'GEMINI_API_KEY_5'}
    updates = {}
    for k, v in body.items():
        if k not in allowed:
            continue
        v = (v or '').strip()
        updates[k] = v
        if v:
            os.environ[k] = v
        else:
            os.environ.pop(k, None)
    if not updates:
        return jsonify({'error': 'No valid keys in request'}), 400
    try:
        _write_env_file(updates)
    except Exception as e:
        return jsonify({'error': f'Failed to save .env: {e}'}), 500
    n = _reload_gemini_keys()
    return jsonify({'success': True, 'keyCount': n, 'configured': bool(n)})

@app.route('/api/config/test', methods=['POST'])
def api_config_test_key():
    """Validate a single Gemini key by calling the embeddings API."""
    body = request.get_json(silent=True) or {}
    key = (body.get('key') or '').strip()
    if not key:
        return jsonify({'ok': False, 'error': 'API key is required'}), 400
    try:
        url = f'https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_EMBED_MODEL}:embedContent?key={key}'
        r = http_req.post(url, json={
            'model': f'models/{GEMINI_EMBED_MODEL}',
            'content': {'parts': [{'text': 'ping'}]}
        }, timeout=10)
        if r.status_code == 200:
            return jsonify({'ok': True, 'message': 'Key is valid'})
        try:
            err = r.json().get('error', {}).get('message', r.text[:200])
        except Exception:
            err = r.text[:200]
        return jsonify({'ok': False, 'status': r.status_code, 'error': err}), 200
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 200


# ─── Document Wiki (auto-generated per upload) ───────

WIKI_DIR = os.path.join(UPLOAD_FOLDER, '_wiki')
os.makedirs(WIKI_DIR, exist_ok=True)

def _doc_corpus_for_wiki(chunks, max_chars=18000):
    """Concatenate enriched chunks (summary + text) into a bounded corpus for one LLM call."""
    parts = []
    used = 0
    for i, c in enumerate(chunks):
        s = (c.get('summary') or '').strip()
        t = (c.get('text') or '').strip()
        sec = (c.get('section') or '').strip()
        head = f'[Chunk {i}'
        if sec: head += f' • {sec}'
        head += ']'
        body = (s + '\n' if s else '') + t
        # truncate per-chunk if huge
        if len(body) > 1200:
            body = body[:1200] + '...'
        block = head + '\n' + body + '\n\n'
        if used + len(block) > max_chars:
            break
        parts.append(block)
        used += len(block)
    return ''.join(parts)

def _build_doc_wiki(doc_id, doc_name, chunks):
    """Generate a structured wiki page for one document via single Gemini call.
    Returns dict {summary, tags, entities, sections, suggested_questions, doc_type, language}.
    Persists to uploads/_wiki/<doc_id>.json + .md.
    """
    if not GEMINI_KEYS or not chunks:
        return None
    corpus = _doc_corpus_for_wiki(chunks)
    if len(corpus) < 50:
        return None
    prompt = f"""You are a librarian. Read this document and produce a STRUCTURED WIKI PAGE in JSON.
Reply in the SAME LANGUAGE as the document (Thai if Thai, English if English).

DOCUMENT NAME: {doc_name}
TOTAL CHUNKS: {len(chunks)}

DOCUMENT CONTENT (excerpts from chunks):
\"\"\"
{corpus}
\"\"\"

Reply with ONLY a JSON object (no markdown fence) with these exact keys:
{{
  "summary": "2-4 sentence overview of what this document is about",
  "doc_type": "one of: report, contract, policy, manual, financial, presentation, article, data, letter, other",
  "language": "th or en or mixed",
  "tags": ["5-10 short topical tags, lowercase, single or two words"],
  "entities": {{
    "people": ["names of people mentioned (max 10)"],
    "organizations": ["companies/depts (max 10)"],
    "places": ["locations (max 10)"],
    "dates": ["important dates/periods (max 10)"],
    "numbers": ["important figures with brief context, e.g. '5.2M baht revenue' (max 10)"]
  }},
  "sections": [
    {{"title": "Section heading", "summary": "1-2 sentences"}}
  ],
  "key_points": ["5-8 bullet points of the most important takeaways"],
  "suggested_questions": ["6-10 useful questions a reader might ask about this document"]
}}

Rules:
- Tags must be useful for cross-document search (e.g. "budget 2024", "hr policy", "vendor contract").
- Skip entity arrays if empty (use []).
- Sections should reflect the actual document structure (max 8).
- Output JSON ONLY, no surrounding text.
"""
    try:
        resp = _gemini_call(prompt, is_embed=False)
        if not resp:
            return None
        # Strip code fence if any
        txt = resp.strip()
        if txt.startswith('```'):
            txt = re.sub(r'^```[a-z]*\n', '', txt)
            txt = re.sub(r'\n```$', '', txt)
        # Find first { ... last }
        i, j = txt.find('{'), txt.rfind('}')
        if i < 0 or j < 0:
            return None
        wiki = json.loads(txt[i:j+1])
    except Exception as e:
        print(f'  ⚠️ wiki gen failed for {doc_id}: {e}', flush=True)
        return None

    wiki['doc_id'] = doc_id
    wiki['doc_name'] = doc_name
    wiki['chunk_count'] = len(chunks)
    wiki['generated_at'] = datetime.utcnow().isoformat() + 'Z'

    # Persist JSON
    try:
        with open(os.path.join(WIKI_DIR, f'{doc_id}.json'), 'w', encoding='utf-8') as f:
            json.dump(wiki, f, ensure_ascii=False, indent=2)
        # Render markdown
        md = _render_wiki_markdown(wiki)
        with open(os.path.join(WIKI_DIR, f'{doc_id}.md'), 'w', encoding='utf-8') as f:
            f.write(md)
    except Exception as e:
        print(f'  ⚠️ wiki save failed for {doc_id}: {e}', flush=True)

    print(f'  📚 wiki built for {doc_name}: {len(wiki.get("tags",[]))} tags, {len(wiki.get("sections",[]))} sections', flush=True)
    return wiki

def _render_wiki_markdown(w):
    """Render wiki dict as a human-readable .md file."""
    lines = []
    lines.append(f"# 📄 {w.get('doc_name','(untitled)')}")
    lines.append('')
    meta = []
    if w.get('doc_type'): meta.append(f"**Type**: {w['doc_type']}")
    if w.get('language'): meta.append(f"**Language**: {w['language']}")
    meta.append(f"**Chunks**: {w.get('chunk_count',0)}")
    if w.get('generated_at'): meta.append(f"**Built**: {w['generated_at'][:10]}")
    lines.append(' • '.join(meta))
    lines.append('')
    if w.get('tags'):
        lines.append('**Tags**: ' + ' '.join(f'`{t}`' for t in w['tags']))
        lines.append('')
    if w.get('summary'):
        lines.append('## 📝 Summary')
        lines.append(w['summary'])
        lines.append('')
    if w.get('key_points'):
        lines.append('## 🎯 Key Points')
        for p in w['key_points']:
            lines.append(f'- {p}')
        lines.append('')
    if w.get('sections'):
        lines.append('## 📑 Sections')
        for s in w['sections']:
            lines.append(f"### {s.get('title','')}")
            lines.append(s.get('summary',''))
            lines.append('')
    ent = w.get('entities') or {}
    if any(ent.get(k) for k in ['people','organizations','places','dates','numbers']):
        lines.append('## 🏷️ Entities')
        for k, label in [('people','People'),('organizations','Organizations'),('places','Places'),('dates','Dates'),('numbers','Key Numbers')]:
            v = ent.get(k) or []
            if v:
                lines.append(f"- **{label}**: " + ', '.join(str(x) for x in v))
        lines.append('')
    if w.get('suggested_questions'):
        lines.append('## ❓ Suggested Questions')
        for q in w['suggested_questions']:
            lines.append(f'- {q}')
        lines.append('')
    return '\n'.join(lines)

def _load_doc_wiki(doc_id):
    p = os.path.join(WIKI_DIR, f'{doc_id}.json')
    if not os.path.exists(p):
        return None
    try:
        with open(p, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return None

def _delete_doc_wiki(doc_id):
    for ext in ('.json', '.md'):
        p = os.path.join(WIKI_DIR, f'{doc_id}{ext}')
        if os.path.exists(p):
            try: os.remove(p)
            except: pass

def _rebuild_wiki_index():
    """Aggregate all per-doc wiki files into uploads/_wiki/index.json.
    Includes: docs list, tag frequency, tag co-occurrence pairs, entity index.
    """
    docs = []
    tag_count = {}
    tag_pairs = {}
    entity_index = {'people': {}, 'organizations': {}, 'places': {}}
    for fn in os.listdir(WIKI_DIR):
        if not fn.endswith('.json') or fn == 'index.json':
            continue
        try:
            with open(os.path.join(WIKI_DIR, fn), 'r', encoding='utf-8') as f:
                w = json.load(f)
        except Exception:
            continue
        did = w.get('doc_id') or fn[:-5]
        # Verify doc still exists in metadata
        if not any(d['id'] == did for d in documents):
            continue
        tags = w.get('tags') or []
        for t in tags:
            tag_count[t] = tag_count.get(t, 0) + 1
        for i, t1 in enumerate(tags):
            for t2 in tags[i+1:]:
                key = tuple(sorted([t1, t2]))
                tag_pairs[key] = tag_pairs.get(key, 0) + 1
        ent = w.get('entities') or {}
        for cat in ['people', 'organizations', 'places']:
            for e in (ent.get(cat) or []):
                entity_index[cat].setdefault(e, []).append(did)
        docs.append({
            'doc_id': did,
            'doc_name': w.get('doc_name', ''),
            'doc_type': w.get('doc_type', ''),
            'language': w.get('language', ''),
            'summary': w.get('summary', ''),
            'tags': tags,
            'chunk_count': w.get('chunk_count', 0),
            'generated_at': w.get('generated_at', ''),
        })
    docs.sort(key=lambda d: d.get('generated_at', ''), reverse=True)
    index = {
        'docs': docs,
        'doc_count': len(docs),
        'tags': sorted(
            [{'tag': k, 'count': v} for k, v in tag_count.items()],
            key=lambda x: -x['count']
        ),
        'tag_pairs': [
            {'a': a, 'b': b, 'count': c}
            for (a, b), c in sorted(tag_pairs.items(), key=lambda x: -x[1])[:40]
        ],
        'entities': {
            cat: [{'name': n, 'docs': ds} for n, ds in sorted(d.items(), key=lambda x: -len(x[1]))[:30]]
            for cat, d in entity_index.items()
        },
        'updated_at': datetime.utcnow().isoformat() + 'Z',
    }
    with open(os.path.join(WIKI_DIR, 'index.json'), 'w', encoding='utf-8') as f:
        json.dump(index, f, ensure_ascii=False, indent=2)
    return index

@app.route('/api/wiki/index')
def api_wiki_index():
    """Return the cross-document wiki index, filtered to docs visible to the current session."""
    pth = os.path.join(WIKI_DIR, 'index.json')
    if not os.path.exists(pth):
        idx = _rebuild_wiki_index()
    else:
        try:
            with open(pth, 'r', encoding='utf-8') as f:
                idx = json.load(f)
        except Exception:
            idx = _rebuild_wiki_index()
    # Filter docs by visible owner
    visible_ids = _visible_doc_ids()
    filtered_docs = [d for d in idx.get('docs', []) if d.get('doc_id') in visible_ids]
    idx = dict(idx)
    idx['docs'] = filtered_docs
    idx['doc_count'] = len(filtered_docs)
    return jsonify(idx)

@app.route('/api/wiki/<doc_id>')
def api_wiki_get(doc_id):
    if doc_id not in _visible_doc_ids():
        return jsonify({'error': 'Document not visible to this session'}), 403
    w = _load_doc_wiki(doc_id)
    if not w:
        return jsonify({'error': 'Wiki not found for this document'}), 404
    # Include rendered markdown for convenience
    w['markdown'] = _render_wiki_markdown(w)
    return jsonify(w)

@app.route('/api/wiki/rebuild/<doc_id>', methods=['POST'])
def api_wiki_rebuild(doc_id):
    """Rebuild wiki for a single doc (uses cached enriched chunks)."""
    if doc_id not in _visible_doc_ids():
        return jsonify({'error': 'Document not visible to this session'}), 403
    cache = _load_enriched_cache(doc_id)
    if not cache or not cache.get('chunks'):
        return jsonify({'error': 'No cached chunks for this document'}), 404
    name = next((d['name'] for d in documents if d['id'] == doc_id), doc_id)
    w = _build_doc_wiki(doc_id, name, cache['chunks'])
    if not w:
        return jsonify({'error': 'Wiki generation failed (no Gemini key or LLM error)'}), 500
    _rebuild_wiki_index()
    return jsonify({'success': True, 'wiki': w})

@app.route('/api/wiki/rebuild', methods=['POST'])
def api_wiki_rebuild_all():
    """Rebuild wikis for ALL documents (background thread)."""
    if not GEMINI_KEYS:
        return jsonify({'error': 'No Gemini API key configured'}), 400
    def _worker():
        n = 0
        for d in list(documents):
            cache = _load_enriched_cache(d['id'])
            if not cache or not cache.get('chunks'):
                continue
            try:
                _build_doc_wiki(d['id'], d['name'], cache['chunks'])
                n += 1
            except Exception as e:
                print(f'  ⚠️ wiki rebuild error for {d["id"]}: {e}', flush=True)
        _rebuild_wiki_index()
        print(f'  📚 wiki rebuild done: {n} docs', flush=True)
    threading.Thread(target=_worker, daemon=True).start()
    return jsonify({'success': True, 'queued': len(documents)})


# ─── Stats ────────────────────────────────────────────

@app.route('/api/stats')
def api_stats():
    # Per-session: only count docs visible to this session (own + public)
    vis = _visible_docs()
    vis_ids = {d['id'] for d in vis}
    enr_info = {}
    for did, data in enriched_store.items():
        if did not in vis_ids: continue
        cks = data.get('chunks', [])
        enr_info[did] = {
            'chunks': len(cks),
            'summaries': sum(1 for c in cks if c.get('summary')),
            'embeddings': sum(1 for c in cks if c.get('embedding') and len(c['embedding']) > 0),
        }
    return jsonify({
        'documentCount': len(vis),
        'totalWords': sum(d.get('wordCount', 0) for d in vis),
        'totalChunks': sum(d.get('chunkCount', len(d.get('chunks', []))) for d in vis),
        'hasGemini': bool(GEMINI_KEYS),
        'geminiKeys': len(GEMINI_KEYS),
        'chatCount': len(chat_history),
        'enrichment': enr_info,
        'dbChunks': _db_total_chunks(),
        'faissVectors': faiss_index.ntotal if faiss_index else 0,
    })


# ═══════════════════════════════════════════════════════
#  Main
# ═══════════════════════════════════════════════════════

# ─── Session / About / Download ───────────────────────
@app.route('/api/session')
def api_session():
    """Return the current session info."""
    sid = getattr(g, 'sid', '')
    n_mine = sum(1 for d in documents if d.get('owner') == sid)
    n_public = sum(1 for d in documents if d.get('owner','public') == 'public')
    return jsonify({
        'sessionId': sid[:8] if sid else '',
        'fullSessionId': sid,
        'myDocs': n_mine,
        'publicDocs': n_public,
        'totalVisible': n_mine + n_public,
    })

@app.route('/api/session/reset', methods=['POST'])
def api_session_reset():
    """Issue a fresh session id (forgets all current uploads from this user)."""
    new = secrets.token_hex(16)
    g._new_sid = new
    g.sid = new
    return jsonify({'success': True, 'sessionId': new[:8]})

@app.route('/about')
def about_page():
    return send_from_directory(os.path.dirname(os.path.abspath(__file__)), 'about.html')

@app.route('/api/download/source')
def api_download_source():
    """Stream a zip of the source code (excluding uploads, caches, backups)."""
    EXCLUDE_DIRS = {'__pycache__', '.git', 'uploads', '.venv', 'venv', 'node_modules', '.wiki/_cache', '.idea', '.vscode'}
    EXCLUDE_PAT = ('.bak', '.bak_', '.pyc', '.pyo', '.log', '.sqlite', '.db')
    # Exact filenames that may contain secrets (allow .env.example through)
    EXCLUDE_FILES = {'.env', '.env.local', '.env.production', '.env.development',
                     '_api_keys.json', 'secrets.json', 'credentials.json',
                     'id_rsa', 'id_dsa', 'id_ecdsa', 'id_ed25519'}
    EXCLUDE_EXT  = ('.key', '.pem', '.pfx', '.p12', '.crt', '.cer')
    # Regex for real Google API keys (AIzaSy + 33 chars). If found inside a file, skip it.
    import re as _re
    KEY_RE = _re.compile(rb'AIzaSy[A-Za-z0-9_\-]{33}')
    buf = io.BytesIO()
    base = os.path.dirname(os.path.abspath(__file__))
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(base):
            dirs[:] = [d for d in dirs if d not in EXCLUDE_DIRS]
            for fn in files:
                if fn in EXCLUDE_FILES: continue
                # Block .env* but keep .env.example
                if fn.startswith('.env') and fn != '.env.example': continue
                if any(p in fn for p in EXCLUDE_PAT): continue
                if any(fn.lower().endswith(e) for e in EXCLUDE_EXT): continue
                full = os.path.join(root, fn)
                arc = os.path.relpath(full, base)
                try:
                    if os.path.getsize(full) > 5*1024*1024: continue
                    # Last-line defense: scan content for real API keys
                    with open(full, 'rb') as rf:
                        data = rf.read()
                    if KEY_RE.search(data):
                        print(f'[download] BLOCKED (contains API key): {arc}', flush=True)
                        continue
                    zf.writestr(arc, data)
                except Exception:
                    continue
    buf.seek(0)
    return send_file(buf, mimetype='application/zip', as_attachment=True,
                     download_name='ai-chat-demo.zip')


# ===== Module-level initialization (runs under both gunicorn and python server.py) =====
_init_database()
_init_faiss_index()
_load_metadata()
_migrate_to_db()
print(f'  ✅ Module init complete: FAISS={faiss_index.ntotal if faiss_index else 0} vecs, keys={len(GEMINI_KEYS)}', flush=True)


if __name__ == '__main__':
    print()
    print('  ╔══════════════════════════════════════════════╗', flush=True)
    print('  ║  🤖 AI Chat Demo — Advanced RAG v2           ║', flush=True)
    print('  ║  Hybrid Search + Enrichment + Embedding      ║', flush=True)
    print('  ╚══════════════════════════════════════════════╝', flush=True)
    print(f'  📂 Uploads : {UPLOAD_FOLDER}', flush=True)
    print(f'  🔑 Gemini  : {len(GEMINI_KEYS)} keys loaded', flush=True)
    for i, k in enumerate(GEMINI_KEYS):
        print(f'     Key {i+1}: {k[:12]}...', flush=True)
    print(f'  🧠 Content : {GEMINI_CONTENT_MODEL}', flush=True)
    print(f'  👁️ OCR     : {GEMINI_OCR_MODEL}', flush=True)
    print(f'  📐 Embed   : {GEMINI_EMBED_MODEL}', flush=True)
    print(f'  📊 Enriched: {len(enriched_store)} docs in store', flush=True)
    for did, data in enriched_store.items():
        cks = data.get('chunks',[])
        ns = sum(1 for c in cks if c.get('summary'))
        ne = sum(1 for c in cks if c.get('embedding') and len(c['embedding']) > 0)
        print(f'     {did}: {len(cks)} chunks, {ns} summaries, {ne} embeddings', flush=True)
    print(f'  💾 SQLite  : {DB_PATH}', flush=True)
    print(f'  🔍 FAISS   : {faiss_index.ntotal if faiss_index else 0:,} vectors', flush=True)
    print(f'  📦 DB chunks: {_db_total_chunks():,}', flush=True)
    port = int(os.environ.get('PORT', 5000))
    print(f'  🌐 URL     : http://0.0.0.0:{port}', flush=True)
    print()
    app.run(host='0.0.0.0', port=port, debug=False)
